import os
import re
import json
import logging  # Add logging import
from flask import (
    Blueprint, render_template, g, request, redirect, url_for, session, flash, jsonify, current_app
)
from datetime import datetime, timezone, date, timedelta # Ensure datetime is available
import traceback # Added traceback import
from werkzeug.utils import secure_filename
from google.cloud.firestore_v1.base_query import FieldFilter
from google.cloud import firestore # Added firestore
from concierge.api.utils import generate_qna_with_gemini
from concierge.utils.gemini_config import genai_enabled, gemini_model
from concierge.utils.file_helpers import allowed_file

# Import config settings for lambda
from concierge.config import LAMBDA_CLIENT

# Import login_required decorator from auth utils
from concierge.auth.utils import login_required

# Import firebase_admin
from firebase_admin import firestore, auth, credentials
from google.cloud.firestore_v1.transforms import ArrayUnion, ArrayRemove

# Import local utils - Use absolute imports from project root
from concierge.utils.ai_helpers import get_relevant_context, format_prompt_with_rag, process_query_with_rag # Absolute import
from concierge.utils.firestore_item_check import check_item_in_firestore, check_item_in_firestore_with_content

# Import date utilities for consistent date handling
from concierge.utils.date_utils import (
    format_date_for_display,
    normalize_reservation_dates,
    to_date_only,
    ensure_date_only_format
)

# Import Firebase setup functions from lambda_src
from concierge.lambda_src.firebase_admin_config import initialize_firebase, get_firestore_client

# Import Firestore client functions
from concierge.utils.firestore_client import list_properties_by_host, create_property, get_property, update_property, list_knowledge_sources, create_knowledge_source, get_user, update_user, list_reservations_by_phone, create_reservation, get_reservation, update_user_pin, verify_user_pin

# Setup logger
logger = logging.getLogger(__name__)

# Create Blueprint
# Note: No url_prefix, these are top-level views. Template folder defaults to project 'templates/'
views_bp = Blueprint('views', __name__, template_folder='../templates')

# Note: The root '/' route is now handled in app.py to direct to either
# 'auth.login' or 'views.dashboard' based on session state.
# We don't need a separate '/' route here unless we want a dedicated landing page
# separate from login.

@views_bp.route('/dashboard')
@login_required # Use the decorator from auth.utils
def dashboard():
    """Main dashboard route - redirects to appropriate dashboard based on user roles"""
    from datetime import datetime, timedelta
    from concierge.utils.firestore_client import get_user
    from concierge.utils.role_helpers import can_access_host_dashboard, get_default_dashboard_path

    user_id = session.get('user_id')
    if not user_id:
        current_app.logger.error("user_id not found in session in dashboard route.")
        return redirect(url_for('auth.login'))

    # Get user data from Firestore
    logger.info(f"Dashboard route: Attempting to get user {user_id} from Firestore...")
    user_data = get_user(user_id)
    if not user_data:
        current_app.logger.error(f"User {user_id} not found in Firestore.")
        
        # Add debugging: Try to check if Firestore is initialized properly
        try:
            from concierge.utils.firestore_client import get_firestore_client
            db = get_firestore_client()
            if db is None:
                current_app.logger.error("Firestore client is None - initialization failed.")
            else:
                current_app.logger.info("Firestore client is initialized.")
        except Exception as init_error:
            current_app.logger.error(f"Error checking Firestore initialization: {init_error}")
        
        # Fallback error handling
        flash('Unable to load user data. Please try logging in again.', 'error')
        return redirect(url_for('auth.logout'))

    # Check if user can access host dashboard
    if not can_access_host_dashboard(user_data):
        # User doesn't have host/property_manager role, redirect to guest dashboard
        return redirect(url_for('views.guest_dashboard'))

    # User has host access, proceed with host dashboard
    return render_host_dashboard(user_id, user_data)

@views_bp.route('/guest')
@login_required
def guest_dashboard():
    """Guest dashboard route"""
    from datetime import datetime, timedelta
    from concierge.utils.firestore_client import get_user, update_user
    from concierge.utils.role_helpers import ensure_guest_role, normalize_user_roles
    from concierge.auth.utils import update_session_roles

    user_id = session.get('user_id')
    if not user_id:
        current_app.logger.error("user_id not found in session in guest dashboard route.")
        return redirect(url_for('auth.login'))

    # Get user data from Firestore
    logger.info(f"Guest dashboard route: Attempting to get user {user_id} from Firestore...")
    user_data = get_user(user_id)
    if not user_data:
        current_app.logger.error(f"User {user_id} not found in Firestore.")
        flash('Unable to load user data. Please try logging in again.', 'error')
        return redirect(url_for('auth.logout'))

    # Ensure user has guest role (silently add if they don't)
    current_roles = normalize_user_roles(user_data)
    if 'guest' not in current_roles:
        logger.info(f"Adding guest role to user {user_id} who accessed guest dashboard")
        updated_user_data = ensure_guest_role(user_data)
        
        # Update in Firestore
        update_success = update_user(user_id, {'role': updated_user_data['role']})
        if update_success:
            user_data = updated_user_data
            # Update session with new roles
            update_session_roles(user_data)
            logger.info(f"Successfully added guest role to user {user_id}")
        else:
            logger.error(f"Failed to update user {user_id} with guest role")

    # Render guest dashboard
    return render_guest_dashboard(user_id, user_data)

def render_host_dashboard(user_id, user_data):
    """Render the host dashboard with host-specific data"""
    from concierge.utils.firestore_client import list_properties_by_host
    from concierge.utils.role_helpers import get_primary_role

    # Get user phone number for context
    phone_number = user_data.get('phoneNumber', user_data.get('PhoneNumber', ''))

    # Get AWS API Gateway WebSocket URL from environment for the dashboard
    websocket_api_url = current_app.config.get('WEBSOCKET_API_URL', 'wss://5qo5m7huzb.execute-api.us-east-1.amazonaws.com/prod/')

    # Hardcoded AWS API Gateway URL
    logger.info("Using hardcoded AWS API Gateway WebSocket URL for host dashboard.")

    # Also set the general websocket URL to the same value
    websocket_url = websocket_api_url

    # Log both URLs for debugging
    logger.info(f"WebSocket URLs for host dashboard: API_URL={websocket_api_url}, URL={websocket_url}")

    # Create initial context for the template
    from datetime import datetime
    context = {
        'user_id': user_id,
        'websocket_url': websocket_url,
        'websocket_api_url': websocket_api_url,  # Add the API URL separately
        'current_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'user_data': user_data,
        'phone_number': phone_number  # Add phone_number to context for all users
    }

    # Get the primary role
    role = get_primary_role(user_data)

    # Get properties where user is listed as host or manager
    properties = list_properties_by_host(user_id)

    context['properties'] = properties

    # Default to first property if available
    if properties:
        context['active_property'] = properties[0]

    # Add user_role directly to context for the host dashboard
    context['user_role'] = role

    # Add display_name and email directly to context for the host dashboard
    context['display_name'] = user_data.get('displayName', user_data.get('DisplayName', 'Host'))
    context['email'] = user_data.get('email', user_data.get('Email', 'Not provided'))

    return render_template('host_dashboard.html', **context)

def render_guest_dashboard(user_id, user_data):
    """Render the guest dashboard with guest-specific data"""
    from concierge.utils.firestore_client import list_reservations_by_phone
    from concierge.utils.phone_utils import get_last_4_digits

    # Get user phone number for context
    phone_number = user_data.get('phoneNumber', user_data.get('PhoneNumber', ''))

    # Get AWS API Gateway WebSocket URL from environment for the dashboard
    websocket_api_url = current_app.config.get('WEBSOCKET_API_URL', 'wss://5qo5m7huzb.execute-api.us-east-1.amazonaws.com/prod/')

    # Hardcoded AWS API Gateway URL
    logger.info("Using hardcoded AWS API Gateway WebSocket URL for guest dashboard.")

    # Also set the general websocket URL to the same value
    websocket_url = websocket_api_url

    # Log both URLs for debugging
    logger.info(f"WebSocket URLs for guest dashboard: API_URL={websocket_api_url}, URL={websocket_url}")

    # Create initial context for the template
    from datetime import datetime
    context = {
        'user_id': user_id,
        'websocket_url': websocket_url,
        'websocket_api_url': websocket_api_url,  # Add the API URL separately
        'current_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'user_data': user_data,
        'phone_number': phone_number  # Add phone_number to context for all users
    }

    # For guests, we'll just pass their phone number and let JavaScript handle fetching reservations
    # This is more robust as it allows the client-side code to handle different data formats

    # Get display name for the guest - prioritize user profile over reservation
    display_name = user_data.get('displayName') or user_data.get('DisplayName') or user_data.get('name')

    # Try to get reservations associated with this user
    guest_name_from_reservation = None
    found_reservation = False

    # Check user's reservationIds field first (like the API does)
    reservations = []
    if user_data and 'reservationIds' in user_data and user_data['reservationIds']:
        logger.info(f"Checking user's reservationIds: {user_data['reservationIds']}")
        from concierge.utils.firestore_client import get_reservation
        for reservation_id in user_data['reservationIds']:
            try:
                reservation_data = get_reservation(reservation_id)
                if reservation_data:
                    reservations.append(reservation_data)
                    logger.info(f"Added reservation from reservationIds: {reservation_id}")
            except Exception as e:
                logger.error(f"Error fetching reservation {reservation_id}: {e}")

    # If no reservations from reservationIds, try phone number lookup
    if not reservations and phone_number:
        logger.info(f"Looking up reservations for phone number: {phone_number}")
        
        try:
            phone_reservations = list_reservations_by_phone(phone_number)
            logger.info(f"Found {len(phone_reservations)} reservations for phone {phone_number}")
            reservations.extend(phone_reservations)
                
        except Exception as e:
            logger.error(f"Error fetching reservations for phone {phone_number}: {e}")
            import traceback
            traceback.print_exc()

    # Process found reservations
    if reservations:
        # Use the most recent reservation as the primary one
        context['reservations'] = reservations
        primary_reservation = reservations[0]  # Assuming sorted by date
        context['primary_reservation'] = primary_reservation
        
        # Extract guest name from reservation
        guest_name_from_reservation = primary_reservation.get('guestName', primary_reservation.get('guest_name'))
        
        logger.info(f"Primary reservation: {primary_reservation.get('id', 'unknown')} for guest: {guest_name_from_reservation}")
        found_reservation = True
    else:
        logger.info(f"No reservations found for user {user_id}")

    if not found_reservation:
        # Add a message to the context if no reservation was found
        logger.warning("No reservation found for guest, the dashboard will show empty state")
        context['no_reservation_message'] = "No active reservations found. Please contact your host for assistance."

    # Determine guest name with proper priority:
    # 1. User profile displayName (highest priority)
    # 2. Guest name from reservation (fallback)
    # 3. "Guest" (final fallback)
    if display_name and display_name.strip() and display_name.strip() != 'Guest':
        context['guest_name'] = display_name.strip()
        context['guest_name_source'] = 'user-profile'
        logger.info(f"Using guest name from user profile: {context['guest_name']}")
    elif found_reservation and guest_name_from_reservation and guest_name_from_reservation.strip() and guest_name_from_reservation.strip() != 'Guest':
        context['guest_name'] = guest_name_from_reservation.strip()
        context['guest_name_source'] = 'reservation'
        logger.info(f"Using guest name from reservation: {context['guest_name']}")
    else:
        context['guest_name'] = 'Guest'
        context['guest_name_source'] = 'default'
        logger.info("Using default guest name: Guest")

    # Check if user has a default PIN (for navbar warning indicator)
    user_pin = user_data.get('pinCode')
    has_default_pin = False
    
    if user_pin and phone_number:
        expected_default_pin = get_last_4_digits(phone_number)
        has_default_pin = (user_pin == expected_default_pin)
        logger.info(f"User {user_id} PIN check: has PIN={bool(user_pin)}, matches default={has_default_pin}")
    
    context['user_has_default_pin'] = has_default_pin
    context['is_temporary_user'] = user_data.get('isTemporary', False)

    return render_template('guest_dashboard.html', **context)

# --- Host Profile Management ---
@views_bp.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    user_id = getattr(g, 'user_id', None)

    # Import Firestore functions
    from concierge.utils.firestore_client import update_user, get_user

    if request.method == 'POST':
        # POST: Update user profile in Firestore
        try:
            display_name = request.form.get('displayName')
            phone_number = request.form.get('phoneNumber')
            email = request.form.get('email')

            # Update data for Firestore (camelCase keys)
            firestore_update_data = {
                'displayName': display_name,
                'phoneNumber': phone_number,
                'email': email
            }

            # Update in Firestore
            success = update_user(user_id, firestore_update_data)

            if success:
                print(f"User profile updated for {user_id}")
                # Redirect with success parameter
                return redirect(url_for('views.profile', updated='true'))
            else:
                print(f"Error updating profile for {user_id}")
                # Render the form again with error message
                user_data = get_user(user_id)
                return render_template('profile.html', user_id=user_id, user_data=user_data,
                                      error="Failed to update profile. Please try again.")

        except Exception as e:
            print(f"Error updating profile for {user_id}: {e}")
            # Render the form again with error message
            user_data = get_user(user_id)
            return render_template('profile.html', user_id=user_id, user_data=user_data,
                                  error=f"Error updating profile: {str(e)}")

    # GET: Display user profile form
    user_data = None
    success = request.args.get('updated') == 'true'

    try:
        # Get user data from Firestore
        user_data = get_user(user_id)
        if not user_data:
            # Should not happen for a logged-in user, but handle defensively
            print(f"Warning: No user document found for logged-in user {user_id}")
            return render_template('profile.html', user_id=user_id, user_data=None,
                                  error="User profile not found. Please contact support.")
    except Exception as e:
        print(f"Error fetching profile for {user_id}: {e}")
        return render_template('profile.html', user_id=user_id, user_data=None,
                              error=f"Error loading profile: {str(e)}")

    print(f"User data for {user_id}: {user_data}")
    return render_template('profile.html', user_id=user_id, user_data=user_data,
                          success="Profile loaded successfully." if success else None)

# --- Profile API Route Removed ---
# This route has been removed to prevent conflicts with the magic link-compatible
# /api/profile route in concierge/api/routes.py which supports both regular and magic link authentication

# --- PIN Management API ---
@views_bp.route('/api/change-pin', methods=['POST'])
@login_required
def change_pin_api():
    """API endpoint for changing user PIN (works for all authenticated users) - Uses OTP verification instead of current PIN"""
    user_id = getattr(g, 'user_id', None)

    try:
        # Get JSON data from request
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        new_pin = data.get('newPin', '').strip()

        if not new_pin:
            return jsonify({"error": "New PIN is required"}), 400

        # Validation
        if len(str(new_pin)) != 4:
            return jsonify({"error": "New PIN must be exactly 4 digits"}), 400

        try:
            # Validate PIN contains only digits
            if not str(new_pin).isdigit():
                return jsonify({"error": "PIN must contain only numbers"}), 400
        except ValueError:
            return jsonify({"error": "PIN must contain only numbers"}), 400
        
        # Update the PIN using the imported function
        # Note: User authentication via OTP verification is already confirmed by @login_required
        success = update_user_pin(user_id, new_pin)
        
        if success:
            logger.info(f"PIN changed successfully for user {user_id} via OTP verification")
            return jsonify({"success": True, "message": "PIN updated successfully"})
        else:
            logger.warning(f"PIN change failed for user {user_id}: Database update failed")
            return jsonify({"success": False, "error": "Failed to update PIN"}), 500

    except Exception as e:
        logger.error(f"Error in PIN change API: {e}")
        return jsonify({"error": "PIN change failed"}), 500

# --- Debug Route for PIN Issues ---
@views_bp.route('/api/debug-user-pin')
@login_required
def debug_user_pin():
    """Debug route to check user PIN data"""
    user_id = getattr(g, 'user_id', None)

    try:
        from concierge.utils.firestore_client import get_firestore_client
        db = get_firestore_client()

        # Get user data
        user_doc = db.collection('users').document(user_id).get()
        if not user_doc.exists:
            return jsonify({"error": "User not found"}), 404

        user_data = user_doc.to_dict()

        # Return debug info (mask sensitive data)
        debug_info = {
            "user_id": user_id,
            "is_temporary": user_id.startswith('temp_magic_'),
            "has_pin": bool(user_data.get('pinCode')),
            "pin_value": user_data.get('pinCode'),  # Show actual PIN for debugging
            "pin_type": type(user_data.get('pinCode')).__name__,
            "phone_number": user_data.get('phoneNumber'),
            "display_name": user_data.get('displayName'),
            "is_temporary_field": user_data.get('isTemporary'),
            "migrated_from": user_data.get('migratedFromTempUser'),
            "pin_set_at": str(user_data.get('pinSetAt')) if user_data.get('pinSetAt') else None
        }

        return jsonify(debug_info)

    except Exception as e:
        logger.error(f"Error in debug user PIN: {e}")
        return jsonify({"error": str(e)}), 500

# --- Property Management ---
@views_bp.route('/properties', methods=['GET'])
@login_required
def properties_list():
    """List properties owned by the host from Firestore"""
    user_id = getattr(g, 'user_id', None)
    properties = []

    try:
        # Get properties for this host from Firestore
        from concierge.utils.firestore_client import list_properties_by_host
        raw_properties = list_properties_by_host(user_id)
        print(f"Fetched {len(raw_properties)} properties for user {user_id}")

        # Log raw properties for debugging
        for i, raw_prop in enumerate(raw_properties):
            print(f"Raw property {i+1}:")
            for key, value in raw_prop.items():
                print(f"  {key}: {value}")

        # Process properties to make them easier to work with in templates
        for prop in raw_properties:
            # Firestore already includes 'id' field
            property_id = prop.get('id', '')
            processed_prop = {
                'id': property_id,
                'name': prop.get('name', 'Unnamed Property'),
                'description': prop.get('description', 'No description'),
                'address': prop.get('address', 'N/A'),
                'raw': prop  # Keep the original data just in case
            }
            properties.append(processed_prop)

        # Add more detailed logging
        for i, prop in enumerate(properties):
            print(f"Debug - Processed Property {i+1}:")
            print(f"  id: {prop['id']}")
            print(f"  name: {prop['name']}")
            print(f"  description: {prop['description']}")
            print(f"  address: {prop['address']}")
    except Exception as e:
        print(f"Error fetching properties from Firestore for user {user_id}: {e}")
        import traceback
        traceback.print_exc()
        # Handle error appropriately, maybe flash message

    return render_template('properties_list.html', user_id=user_id, properties=properties)

@views_bp.route('/properties/new', methods=['GET', 'POST'])
@login_required
def property_new():
    """Create a new property using Firestore"""
    user_id = getattr(g, 'user_id', None)
    if request.method == 'POST':
        # Import Firestore client function
        from concierge.utils.firestore_client import create_property
        from datetime import datetime, timezone
        import traceback

        try:
            prop_name = request.form.get('name')
            prop_address = request.form.get('address')
            prop_description = request.form.get('description')
            prop_ical_url = request.form.get('ical_url') # Get the iCal URL

            # Get check-in and check-out times from form
            check_in_time = request.form.get('check_in_time')
            check_out_time = request.form.get('check_out_time')

            # Get WiFi details from form
            wifi_network = request.form.get('wifi_network')
            wifi_password = request.form.get('wifi_password')

            if not prop_name: # Basic validation
                # flash('Property name is required.', 'warning')
                return render_template('property_edit.html', user_id=user_id, property=request.form, error="Property name is required.")

            # Prepare the property data
            new_prop_data = {
                'hostId': user_id,
                'name': prop_name,
                'address': prop_address,
                'description': prop_description,
                'icalUrl': prop_ical_url, # Add iCal URL
                'checkInTime': check_in_time,
                'checkOutTime': check_out_time,
                'createdAt': datetime.now(timezone.utc),
                'updatedAt': datetime.now(timezone.utc)
            }

            # Add WiFi details if provided
            if wifi_network or wifi_password:
                new_prop_data['wifiDetails'] = {
                    'network': wifi_network,
                    'password': wifi_password
                }

            # Generate a new ID for the property
            import uuid
            property_id = str(uuid.uuid4())

            # Create the property in Firestore
            success = create_property(property_id, new_prop_data)

            if success:
                print(f"Created new property {property_id} for user {user_id}")
                # flash('Property created successfully!', 'success')
                return redirect(url_for('views.properties_list'))
            else:
                print(f"Failed to create property in Firestore for user {user_id}")
                return render_template('property_edit.html', user_id=user_id, property=request.form, error="Failed to create property in Firestore")

        except Exception as e:
            print(f"Error creating property in Firestore for user {user_id}: {e}")
            traceback.print_exc()
            # Render the form again with submitted data and error message
            return render_template('property_edit.html', user_id=user_id, property=request.form, error=str(e))

    # GET: Display form to create a new property
    return render_template('property_edit.html', user_id=user_id, property=None)

@views_bp.route('/properties/<property_id>/edit', methods=['GET', 'POST'])
@login_required
def property_edit(property_id):
    """Edit a property using Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import Firestore client functions
    from concierge.utils.firestore_client import get_property, update_property
    from datetime import datetime, timezone
    import traceback

    property_data = None

    try:
        # --- Verify Ownership & Fetch Data (for both GET and POST validation) ---
        raw_property_data = get_property(property_id)
        if not raw_property_data:
            print(f"Attempt to edit non-existent property {property_id} in Firestore by user {user_id}")
            return redirect(url_for('views.properties_list'))

        # Process property data to make it easier to work with in templates
        property_data = {
            'id': property_id,
            'name': raw_property_data.get('name', 'Unnamed Property'),
            'description': raw_property_data.get('description', ''),
            'address': raw_property_data.get('address', ''),
            'ical_url': raw_property_data.get('icalUrl', ''),
            'wifi_details': raw_property_data.get('wifiDetails', {}),
            'check_in_time': raw_property_data.get('checkInTime', ''),
            'check_out_time': raw_property_data.get('checkOutTime', '')
        }
        print(f"Property data prepared for edit template: {property_data}")

        # Check if user is the owner
        host_id = raw_property_data.get('hostId')
        if host_id != user_id:
            print(f"Unauthorized attempt by user {user_id} to edit property {property_id} owned by {host_id}")
            return redirect(url_for('views.properties_list'))
        # --- End Verification ---

        if request.method == 'POST':
            # POST: Update property details in Firestore
            prop_name = request.form.get('name')
            prop_address = request.form.get('address')
            prop_description = request.form.get('description')
            prop_ical_url = request.form.get('ical_url') # Get the iCal URL

            # Get check-in and check-out times from form
            check_in_time = request.form.get('check_in_time')
            check_out_time = request.form.get('check_out_time')

            # Get WiFi details from form
            wifi_network = request.form.get('wifi_network')
            wifi_password = request.form.get('wifi_password')

            if not prop_name: # Basic validation
                # Re-render with existing data and error
                return render_template('property_edit.html', user_id=user_id, property=property_data, error="Property name is required.")

            update_data = {
                'name': prop_name,
                'address': prop_address,
                'description': prop_description,
                'icalUrl': prop_ical_url,
                'checkInTime': check_in_time,
                'checkOutTime': check_out_time,
                'updatedAt': datetime.now(timezone.utc)
            }

            # Add WiFi details if provided
            if wifi_network or wifi_password:
                update_data['wifiDetails'] = {
                    'network': wifi_network,
                    'password': wifi_password
                }

            # Update in Firestore
            success = update_property(property_id, update_data)

            if success:
                print(f"Updated property {property_id} for user {user_id}")
                return redirect(url_for('views.properties_list'))
            else:
                print(f"Failed to update property {property_id} in Firestore")
                return render_template('property_edit.html', user_id=user_id, property=property_data, error="Failed to update property in Firestore")

    except Exception as e:
        print(f"Error accessing or updating property {property_id} in Firestore for user {user_id}: {e}")
        traceback.print_exc()
        return redirect(url_for('views.properties_list'))

    # GET: Display form to edit the existing property
    # property_data should be populated from the try block above
    return render_template('property_edit.html', user_id=user_id, property=property_data)

@views_bp.route('/properties/<property_id>/delete', methods=['POST'])
@login_required
def property_delete(property_id):
    """Delete a property and related data from Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import Firestore client functions
    from concierge.utils.firestore_client import get_property, delete_property
    import traceback

    try:
        # Verify Ownership
        property_data = get_property(property_id)

        if not property_data:
            print(f"Attempt to delete non-existent property {property_id} in Firestore by user {user_id}")
            return redirect(url_for('views.properties_list'))

        # Check if user is the owner
        host_id = property_data.get('hostId')
        if host_id != user_id:
            print(f"Unauthorized attempt by user {user_id} to delete property {property_id} owned by {host_id}")
            return redirect(url_for('views.properties_list'))

        # Delete the property from Firestore
        success = delete_property(property_id)

        if success:
            print(f"Property {property_id} successfully deleted from Firestore by user {user_id}")
        else:
            print(f"Failed to delete property {property_id} from Firestore")

        return redirect(url_for('views.properties_list'))
    except Exception as e:
        print(f"Error deleting property {property_id} from Firestore: {e}")
        traceback.print_exc()
        return redirect(url_for('views.properties_list'))

# --- Conversation History Management ---
@views_bp.route('/properties/<property_id>/conversations', methods=['GET'])
@login_required
def property_conversations(property_id):
    """Display conversation history for a property"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.dynamodb_client import get_property, list_property_conversations
    from concierge.utils.firestore_client import get_property as get_firestore_property
    from concierge.utils.firestore_client import list_property_reservations as list_reservations_by_property
    from concierge.utils.ai_helpers import generate_conversation_summary
    import traceback

    property_data = None
    conversations = []
    reservations = {}

    try:
        # --- Verify Ownership & Fetch Property Data ---
        # Try DynamoDB first
        raw_property_data = get_property(property_id)

        # If not found in DynamoDB, try Firestore as fallback
        if not raw_property_data:
            raw_property_data = get_firestore_property(property_id)

        if not raw_property_data:
            print(f"Conversation history access attempt for non-existent property {property_id} by user {user_id}")
            flash('Property not found.', 'danger')
            return redirect(url_for('views.properties_list'))

        # Process property data to make it easier to work with in templates
        property_data = {
            'id': property_id,
            'name': raw_property_data.get('name', raw_property_data.get('Name', 'Unnamed Property')),
            'description': raw_property_data.get('description', raw_property_data.get('Description', '')),
            'address': raw_property_data.get('address', raw_property_data.get('Address', '')),
        }

        # Check if user is the owner (handle both Firestore and DynamoDB formats)
        host_id = raw_property_data.get('hostId', raw_property_data.get('HostId'))
        if host_id != user_id:
            print(f"Unauthorized conversation history access attempt by user {user_id} for property {property_id} owned by {host_id}")
            flash('You do not have permission to access this property\'s conversations.', 'danger')
            return redirect(url_for('views.properties_list'))
        # --- End Verification ---

        # Fetch all reservations for this property to link conversations to guests
        try:
            # Try to get reservations from DynamoDB first
            from concierge.utils.dynamodb_client import list_property_reservations
            all_reservations = list_property_reservations(property_id)

            # If no reservations found in DynamoDB, try Firestore as fallback
            if not all_reservations:
                all_reservations = list_reservations_by_property(property_id)

            # Create a lookup dictionary by guest name, phone, and reservation ID
            for reservation in all_reservations:
                # Normalize field names (handle both Firestore and DynamoDB formats)
                reservation_data = {
                    'guestName': reservation.get('guestName', reservation.get('GuestName', '')),
                    'guestPhone': reservation.get('guestPhone', reservation.get('GuestPhoneNumber', '')),
                    'checkInDate': reservation.get('startDate', reservation.get('StartDate', '')),
                    'checkOutDate': reservation.get('endDate', reservation.get('EndDate', '')),
                    'reservationId': reservation.get('reservationId', reservation.get('ReservationId', ''))
                }

                # If no explicit reservation ID, try to extract from SK
                if not reservation_data['reservationId'] and 'SK' in reservation and reservation['SK'].startswith('RESERVATION#'):
                    reservation_data['reservationId'] = reservation['SK'].replace('RESERVATION#', '')

                # Store by guest name (lowercase for case-insensitive matching)
                guest_name = reservation_data['guestName'].lower()
                if guest_name:
                    reservations[guest_name] = reservation_data

                # Store by phone number (last 4 digits and full number)
                phone = reservation_data['guestPhone']
                if phone:
                    # Store by full phone
                    reservations[phone] = reservation_data
                    # Store by last 4 digits
                    if len(phone) >= 4:
                        reservations[phone[-4:]] = reservation_data

                # Store by reservation ID
                if reservation_data['reservationId']:
                    reservations[reservation_data['reservationId']] = reservation_data

            print(f"Loaded {len(all_reservations)} reservations for property {property_id}")
        except Exception as res_err:
            print(f"Error loading reservations: {res_err}")
            traceback.print_exc()

        # Fetch conversations from DynamoDB
        conversations = list_property_conversations(property_id)

        # Process conversations to add summaries if they don't exist
        processed_conversations = []
        for conv in conversations:
            # Add a processed version with additional data
            processed_conv = {
                'id': conv.get('ConversationId', conv.get('SK', '').replace('CONVERSATION#', '')),
                'start_time': conv.get('StartTime'),
                'last_update_time': conv.get('LastUpdateTime'),
                'channel': conv.get('Channel', 'unknown'),
                'guest_name': conv.get('GuestName', 'Unknown Guest'),
                'message_count': conv.get('MessageCount', 0),
                'messages': conv.get('Messages', []),
                'summary': conv.get('Summary', ''),
                'reservation_id': conv.get('ReservationId')
            }

            # First try to find reservation by direct reservation ID
            if processed_conv['reservation_id'] and processed_conv['reservation_id'] in reservations:
                processed_conv['reservation'] = reservations[processed_conv['reservation_id']]
                print(f"Found reservation for conversation by reservation ID: {processed_conv['reservation_id']}")
            else:
                # Try to find matching reservation by guest name
                guest_name = processed_conv['guest_name'].lower()
                if guest_name in reservations:
                    processed_conv['reservation'] = reservations[guest_name]
                    print(f"Found reservation for guest '{guest_name}' by name")
                else:
                    # Try to find by phone number if available
                    guest_phone = None
                    # Look for phone number in messages
                    if processed_conv['messages']:
                        for msg in processed_conv['messages']:
                            if msg.get('role') == 'user' and msg.get('phone_number'):
                                guest_phone = msg.get('phone_number')
                                break

                    # If we found a phone number, try to match it
                    if guest_phone and guest_phone in reservations:
                        processed_conv['reservation'] = reservations[guest_phone]
                        print(f"Found reservation for guest '{guest_name}' by phone {guest_phone}")
                    elif guest_phone and len(guest_phone) >= 4 and guest_phone[-4:] in reservations:
                        # Try matching by last 4 digits
                        processed_conv['reservation'] = reservations[guest_phone[-4:]]
                        print(f"Found reservation for guest '{guest_name}' by last 4 digits {guest_phone[-4:]}")

            # If we found a reservation, update the conversation with the reservation ID if it's missing
            if processed_conv.get('reservation'):
                # Update the guest name if it was "Unknown Guest" or starts with "Guest"
                if ((processed_conv['guest_name'] == 'Unknown Guest' or
                     processed_conv['guest_name'].startswith('Guest ')) and
                    processed_conv['reservation'].get('guestName')):
                    processed_conv['guest_name'] = processed_conv['reservation']['guestName']
                    print(f"Updated guest name from '{processed_conv['guest_name']}' to '{processed_conv['reservation']['guestName']}'")

                # Update the reservation ID in DynamoDB if it's missing
                if not processed_conv['reservation_id'] and processed_conv['reservation'].get('reservationId'):
                    try:
                        from concierge.utils.dynamodb_client import update_conversation
                        update_success = update_conversation(
                            property_id=property_id,
                            conversation_id=processed_conv['id'],
                            update_data={'ReservationId': processed_conv['reservation']['reservationId']}
                        )
                        if update_success:
                            processed_conv['reservation_id'] = processed_conv['reservation']['reservationId']
                            print(f"Updated conversation {processed_conv['id']} with reservation ID {processed_conv['reservation_id']}")
                    except Exception as update_err:
                        print(f"Error updating conversation with reservation ID: {update_err}")
                        traceback.print_exc()

            # Ensure messages have proper timestamps
            if processed_conv['messages']:
                for msg in processed_conv['messages']:
                    if not msg.get('timestamp') and processed_conv.get('start_time'):
                        msg['timestamp'] = processed_conv['start_time']

            # Generate summary if it doesn't exist and there are messages
            if not processed_conv['summary'] and processed_conv['messages']:
                try:
                    # Find matching reservation for context
                    reservation_info = processed_conv.get('reservation')

                    # Generate summary using AI with property and reservation context
                    summary = generate_conversation_summary(
                        processed_conv['messages'],
                        property_context=property_data,
                        reservation_info=reservation_info,
                        guest_name=processed_conv.get('guest_name')  # Pass guest name from conversation
                    )
                    processed_conv['summary'] = summary

                    # Update the conversation in DynamoDB with the summary
                    from concierge.utils.dynamodb_client import update_conversation
                    update_conversation(property_id, processed_conv['id'], {'Summary': summary})
                except Exception as summary_err:
                    print(f"Error generating summary: {summary_err}")
                    traceback.print_exc()
                    processed_conv['summary'] = "No summary available."

            processed_conversations.append(processed_conv)

        # Sort conversations by last update time (newest first)
        processed_conversations.sort(key=lambda x: x.get('last_update_time', ''), reverse=True)

        # Render the template with the processed conversations
        try:
            return render_template('property_conversations.html',
                                user_id=user_id,
                                property=property_data,
                                conversations=processed_conversations)
        except Exception as template_err:
            print(f"Template rendering error: {template_err}")
            traceback.print_exc()
            flash(f'Error rendering conversation history: {template_err}', 'danger')
            return redirect(url_for('views.properties_list'))

    except Exception as e:
        print(f"Error accessing conversation history for property {property_id}: {e}")
        traceback.print_exc()
        flash(f'Error accessing conversation history: {e}', 'danger')
        return redirect(url_for('views.properties_list'))

@views_bp.route('/properties/<property_id>/conversations/<conversation_id>/render')
@login_required
def render_conversation(property_id, conversation_id):
    """Render a conversation using the partial template"""
    user_id = getattr(g, 'user_id', None)
    # Initialize reservation variable to None to avoid UnboundLocalError
    reservation = None

    # Import required functions
    from concierge.utils.firestore_client import get_property
    from concierge.utils.dynamodb_client import list_property_conversations, list_conversations_by_property
    import traceback  # Import traceback module explicitly to avoid UnboundLocalError

    # Verify property ownership
    property_data = get_property(property_id)
    if not property_data or property_data.get('hostId') != user_id:
        return jsonify({"error": "Unauthorized"}), 403

    try:
        # Get the conversation from DynamoDB
        # We need to use a different approach since get_conversation uses a different table
        # Let's use list_property_conversations and filter for the specific conversation

        # Add more detailed logging
        print(f"Fetching conversation: ID={conversation_id}, Property ID={property_id}")

        # Get all conversations for the property from both tables
        all_conversations = list_property_conversations(property_id)

        # Also try the main table using list_conversations_by_property

        try:
            # Try to get conversations from the main table
            main_table_conversations = list_conversations_by_property(property_id)
            if main_table_conversations:
                print(f"Found {len(main_table_conversations)} conversations in the main table")
                all_conversations.extend(main_table_conversations)
        except Exception as e:
            print(f"Error getting conversations from main table: {e}")

        # Find the specific conversation by ID
        conversation = None
        print(f"Searching through {len(all_conversations)} conversations")

        # Print the first conversation's keys for debugging
        if all_conversations and len(all_conversations) > 0:
            print(f"First conversation keys: {list(all_conversations[0].keys())}")
            print(f"First conversation SK: {all_conversations[0].get('SK', 'No SK')}")
            print(f"First conversation ConversationId: {all_conversations[0].get('ConversationId', 'No ConversationId')}")

        for i, conv in enumerate(all_conversations):
            conv_id = conv.get('ConversationId', '')
            if not conv_id and 'SK' in conv:
                # Extract ID from SK if ConversationId is not present
                sk = conv.get('SK', '')
                if sk.startswith('CONVERSATION#'):
                    parts = sk.split('#')
                    if len(parts) > 1:
                        conv_id = parts[-1]
                        print(f"Extracted ID {conv_id} from SK {sk}")

            print(f"Conversation {i}: ID={conv_id}, looking for {conversation_id}")

            if conv_id == conversation_id:
                conversation = conv
                print(f"Found matching conversation: {conv_id}")
                break

        print(f"Found {len(all_conversations)} conversations, looking for ID={conversation_id}")
        if not conversation:
            print(f"ERROR: Conversation not found: ID={conversation_id}, Property ID={property_id}")
            return jsonify({
                "error": "Conversation not found",
                "conversation_id": conversation_id,
                "property_id": property_id
            }), 404

        # Process the conversation data
        print(f"Conversation found: {conversation.get('ConversationId', 'No ID')}")
        print(f"Conversation keys: {list(conversation.keys())}")

        # Get reservation ID from conversation
        reservation_id = conversation.get('ReservationId')
        reservation_data = None

        # If we have a reservation ID, try to get the reservation details using the same approach as property_conversations
        if reservation_id:
            try:
                # First, try direct lookup by reservation ID (NEW APPROACH)
                from concierge.utils.firestore_client import get_reservation
                
                reservation = None
                direct_reservation = get_reservation(reservation_id)
                if direct_reservation:
                    print(f"Found reservation by direct ID lookup: {reservation_id}")
                    reservation = {
                        'guestName': direct_reservation.get('guestName', ''),
                        'guestPhone': direct_reservation.get('guestPhoneNumber', ''),
                        'checkInDate': direct_reservation.get('startDate', ''),
                        'checkOutDate': direct_reservation.get('endDate', ''),
                        'reservationId': direct_reservation.get('id', reservation_id)
                    }
                    print(f"Direct reservation data: {reservation}")
                else:
                    print(f"Direct reservation lookup failed for ID: {reservation_id}")
                    
                    # Fallback to property-scoped lookup (EXISTING APPROACH)
                    # Get all reservations for this property and create lookup dictionary (same as property_conversations)
                    from concierge.utils.dynamodb_client import list_property_reservations
                    from concierge.utils.firestore_client import list_property_reservations as list_reservations_by_property
                    
                    # Try to get reservations from DynamoDB first
                    all_reservations = list_property_reservations(property_id)

                    # If no reservations found in DynamoDB, try Firestore as fallback
                    if not all_reservations:
                        all_reservations = list_reservations_by_property(property_id)

                    # Create a lookup dictionary by reservation ID
                    reservations_lookup = {}
                    for reservation_item in all_reservations:
                        # Normalize field names (handle both Firestore and DynamoDB formats)
                        reservation_data_normalized = {
                            'guestName': reservation_item.get('guestName', reservation_item.get('GuestName', '')),
                            'guestPhone': reservation_item.get('guestPhone', reservation_item.get('GuestPhoneNumber', '')),
                            'checkInDate': reservation_item.get('startDate', reservation_item.get('StartDate', '')),
                            'checkOutDate': reservation_item.get('endDate', reservation_item.get('EndDate', '')),
                            'reservationId': reservation_item.get('reservationId', reservation_item.get('ReservationId', reservation_item.get('id', '')))
                        }

                        # If no explicit reservation ID, try to extract from SK
                        if not reservation_data_normalized['reservationId'] and 'SK' in reservation_item and reservation_item['SK'].startswith('RESERVATION#'):
                            reservation_data_normalized['reservationId'] = reservation_item['SK'].replace('RESERVATION#', '')

                        # Store by reservation ID
                        if reservation_data_normalized['reservationId']:
                            reservations_lookup[reservation_data_normalized['reservationId']] = reservation_data_normalized

                    # Look up the specific reservation
                    if reservation_id in reservations_lookup:
                        reservation = reservations_lookup[reservation_id]
                        print(f"Found reservation by property-scoped ID lookup: {reservation_id}")
                        print(f"Property-scoped reservation data: {reservation}")
                    else:
                        print(f"Reservation {reservation_id} not found in reservations lookup for property {property_id}")
                        print(f"Available reservation IDs in property: {list(reservations_lookup.keys())}")
                    
            except Exception as res_err:
                print(f"Error getting reservation {reservation_id}: {res_err}")
                traceback.print_exc()

        processed_conv = {
            'id': conversation.get('ConversationId', conversation.get('SK', '').replace('CONVERSATION#', '')),
            'start_time': conversation.get('StartTime'),
            'last_update_time': conversation.get('LastUpdateTime'),
            'channel': conversation.get('Channel', 'unknown'),
            'guest_name': conversation.get('GuestName', 'Unknown Guest'),
            'message_count': conversation.get('MessageCount', 0),
            'messages': conversation.get('Messages', []),
            'summary': conversation.get('Summary', ''),
            'reservation': reservation,
            'guest_phone': conversation.get('GuestPhone', '')  # Add phone number from conversation
        }

        # If reservation is not directly attached to the conversation, try to find it
        if not processed_conv['reservation']:
            try:
                # Import required functions
                from concierge.utils.firestore_client import list_property_reservations as list_firestore_reservations
                from concierge.utils.dynamodb_client import list_property_reservations as list_dynamodb_reservations
                import traceback

                # Get all reservations for this property from both sources
                all_reservations = list_firestore_reservations(property_id)

                # Also try DynamoDB
                try:
                    dynamodb_reservations = list_dynamodb_reservations(property_id)
                    if dynamodb_reservations:
                        print(f"Found {len(dynamodb_reservations)} reservations in DynamoDB")
                        all_reservations.extend(dynamodb_reservations)
                except Exception as ddb_err:
                    print(f"Error listing DynamoDB reservations: {ddb_err}")
                    traceback.print_exc()

                print(f"Found total of {len(all_reservations)} reservations for property {property_id}")

                # Try to match by guest name
                guest_name = processed_conv['guest_name'].lower()
                for reservation in all_reservations:
                    # Normalize field names between Firestore and DynamoDB
                    res_guest_name = reservation.get('guestName', reservation.get('GuestName', '')).lower()
                    if res_guest_name and res_guest_name == guest_name:
                        # Create normalized reservation data structure (same as property_conversations)
                        processed_conv['reservation'] = {
                            'guestName': reservation.get('guestName', reservation.get('GuestName', '')),
                            'guestPhone': reservation.get('guestPhone', reservation.get('GuestPhoneNumber', '')),
                            'checkInDate': reservation.get('startDate', reservation.get('StartDate', '')),
                            'checkOutDate': reservation.get('endDate', reservation.get('EndDate', '')),
                            'reservationId': reservation.get('reservationId', reservation.get('ReservationId', ''))
                        }
                        print(f"Found reservation for guest '{guest_name}' by name")
                        print(f"Original reservation by name: {reservation}")
                        print(f"Final reservation object by name: {processed_conv['reservation']}")
                        break

                    # If still no match, try to find by phone number in messages or conversation
                    if not processed_conv['reservation']:
                        guest_phone = processed_conv.get('guest_phone')  # Check conversation phone first

                        # If no phone in conversation, check messages
                        if not guest_phone and processed_conv['messages']:
                            for msg in processed_conv['messages']:
                                if msg.get('role') == 'user' and msg.get('phone_number'):
                                    guest_phone = msg.get('phone_number')
                                    break

                        if guest_phone:
                            for reservation in all_reservations:
                                # Check both field names for the phone number
                                res_phone = reservation.get('guestPhone', reservation.get('GuestPhoneNumber', ''))
                                if res_phone and res_phone == guest_phone:
                                    # Create normalized reservation data structure (same as property_conversations)
                                    processed_conv['reservation'] = {
                                        'guestName': reservation.get('guestName', reservation.get('GuestName', '')),
                                        'guestPhone': res_phone,
                                        'checkInDate': reservation.get('startDate', reservation.get('StartDate', '')),
                                        'checkOutDate': reservation.get('endDate', reservation.get('EndDate', '')),
                                        'reservationId': reservation.get('reservationId', reservation.get('ReservationId', ''))
                                    }
                                    print(f"Found reservation for guest by phone {guest_phone}")
                                    print(f"Original reservation by phone: {reservation}")
                                    print(f"Final reservation object by phone: {processed_conv['reservation']}")
                                    break
                                elif len(guest_phone) >= 4 and res_phone and res_phone.endswith(guest_phone[-4:]):
                                    # Create normalized reservation data structure (same as property_conversations)
                                    processed_conv['reservation'] = {
                                        'guestName': reservation.get('guestName', reservation.get('GuestName', '')),
                                        'guestPhone': res_phone,
                                        'checkInDate': reservation.get('startDate', reservation.get('StartDate', '')),
                                        'checkOutDate': reservation.get('endDate', reservation.get('EndDate', '')),
                                        'reservationId': reservation.get('reservationId', reservation.get('ReservationId', ''))
                                    }
                                    print(f"Found reservation for guest by last 4 digits {guest_phone[-4:]}")
                                    print(f"Original reservation by last 4: {reservation}")
                                    print(f"Final reservation object by last 4: {processed_conv['reservation']}")
                                    break
            except Exception as res_err:
                print(f"Error finding reservation: {res_err}")
                traceback.print_exc()

        print(f"Processed conversation: ID={processed_conv['id']}, Guest={processed_conv['guest_name']}, Messages={processed_conv['message_count']}")

        # Ensure messages have proper timestamps
        if processed_conv['messages']:
            for msg in processed_conv['messages']:
                if not msg.get('timestamp') and processed_conv.get('start_time'):
                    msg['timestamp'] = processed_conv['start_time']

        # Generate summary if it doesn't exist and there are messages
        if not processed_conv['summary'] and processed_conv['messages'] and len(processed_conv['messages']) > 0:
            try:
                # Import required functions
                from concierge.utils.ai_helpers import generate_conversation_summary
                from concierge.utils.dynamodb_client import update_conversation

                # Generate summary using AI with property and reservation context
                summary = generate_conversation_summary(
                    processed_conv['messages'],
                    property_context=property_data,
                    reservation_info=processed_conv.get('reservation'),
                    guest_name=processed_conv.get('guest_name')  # Pass guest name from conversation
                )

                if summary:
                    processed_conv['summary'] = summary

                    # Update the conversation in DynamoDB with the summary
                    try:
                        update_conversation(property_id, processed_conv['id'], {'Summary': summary})
                        print(f"Updated conversation {processed_conv['id']} with new summary")
                    except Exception as update_err:
                        print(f"Error updating conversation with summary: {update_err}")
                        traceback.print_exc()
            except Exception as summary_err:
                print(f"Error generating summary: {summary_err}")
                traceback.print_exc()

        # Use a simple standalone template for testing
        html = render_template('conversation_detail.html', conversation=processed_conv)

        # Log the HTML length for debugging
        print(f"Generated HTML length: {len(html)}")

        return html

    except Exception as e:
        print(f"Error rendering conversation: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# --- Knowledge Base Management (per property) ---
@views_bp.route('/properties/<property_id>/knowledge', methods=['GET'])
@login_required
def knowledge_base(property_id):
    """Display unified knowledge management for a property"""
    user_id = getattr(g, 'user_id', None)
    user_role = session.get('user_role')

    # Import required functions
    from concierge.utils.firestore_client import get_property, list_knowledge_items_by_property
    import traceback

    property_data = None
    knowledge_items = []

    try:
        # --- Verify Ownership & Fetch Property Data --- #
        raw_property_data = get_property(property_id)

        if not raw_property_data:
            print(f"Knowledge base access attempt for non-existent property {property_id} in Firestore by user {user_id}")
            flash('Property not found.', 'danger')
            return redirect(url_for('views.properties_list'))

        # Process property data to make it easier to work with in templates
        property_data = {
            'id': property_id,
            'name': raw_property_data.get('name', 'Unnamed Property'),
            'description': raw_property_data.get('description', ''),
            'address': raw_property_data.get('address', ''),
            'ical_url': raw_property_data.get('icalUrl', ''),
        }
        print(f"Property data prepared for unified knowledge template: {property_data}")

        # Check if user is the owner
        host_id = raw_property_data.get('hostId')
        if host_id != user_id:
            print(f"Unauthorized KB access attempt by user {user_id} for property {property_id} owned by {host_id}")
            flash('You do not have permission to access this knowledge base.', 'danger')
            return redirect(url_for('views.properties_list'))
        # --- End Verification ---

    except Exception as e:
        print(f"Error accessing knowledge base for property {property_id} in Firestore for user {user_id}: {e}")
        traceback.print_exc()
        flash(f'Error accessing knowledge base: {e}', 'danger')
        return redirect(url_for('views.properties_list'))

    # Get host name if available
    host_name = "Your Host"
    host_id = raw_property_data.get('hostId')
    if host_id:
        try:
            # Get the host's user profile
            host_data = get_user(host_id)
            if host_data:
                # Try different possible field names for the host's name
                host_name = host_data.get('name',
                            host_data.get('displayName',
                            host_data.get('fullName', host_id)))
        except Exception as e:
            print(f"Error getting host name for host ID {host_id}: {e}")

    # GET: Display unified knowledge management page
    return render_template('knowledge_unified.html',
                           user_id=user_id,
                           property=property_data,
                           host_name=host_name)

@views_bp.route('/properties/<property_id>/knowledge/upload', methods=['POST'])
@login_required
def knowledge_upload(property_id):
    """Handle knowledge file upload using Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_property, create_knowledge_source, create_knowledge_item, update_knowledge_source
    import traceback
    from datetime import datetime, timezone

    try:
        # --- Verify Ownership (Crucial for POST actions) ---
        property_data = get_property(property_id)

        if not property_data:
            print(f"Unauthorized upload attempt for non-existent property {property_id} in Firestore by user {user_id}")
            flash('Property not found in Firestore.', 'danger')
            return redirect(url_for('views.properties_list'))

        # Check if user is the owner
        host_id = property_data.get('hostId')
        if host_id != user_id:
            print(f"Unauthorized upload attempt for property {property_id} by user {user_id}")
            flash('Permission denied.', 'danger')
            return redirect(url_for('views.properties_list'))
        # --- End Verification ---

        property_details = property_data # Already have the property details
        property_details['id'] = property_id # Add id for convenience

        if 'knowledgeFile' not in request.files:
            flash('No file part in request.', 'warning')
            print("No file part in request")
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        file = request.files['knowledgeFile']
        if file.filename == '':
            flash('No selected file.', 'warning')
            print("No file selected")
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        if file and allowed_file(file.filename):
            # Create a secure filename and save the file
            filename = secure_filename(file.filename)
            upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], property_id)
            os.makedirs(upload_dir, exist_ok=True)
            file_path = os.path.join(upload_dir, filename)
            file.save(file_path)
            print(f"Saved file to {file_path}")

            # Extract text from the file
            import magic  # python-magic for MIME type detection
            # Detect MIME type for more reliable file type detection
            mime = magic.Magic(mime=True)
            mime_type = mime.from_file(file_path)
            print(f"Detected MIME type: {mime_type}")

            # Process based on file type
            extracted_text = ""
            if mime_type == 'text/plain':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()
            elif mime_type == 'application/pdf':
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        extracted_text += page_text + "\n\n"
                except Exception as pdf_err:
                    print(f"Error extracting PDF text: {pdf_err}")
                    flash('Error processing PDF file.', 'danger')
                    os.remove(file_path)  # Clean up file
                    return redirect(url_for('views.knowledge_base', property_id=property_id))
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    import docx
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        extracted_text += para.text + "\n"
                except Exception as docx_err:
                    print(f"Error extracting DOCX text: {docx_err}")
                    flash('Error processing Word document.', 'danger')
                    os.remove(file_path)  # Clean up file
                    return redirect(url_for('views.knowledge_base', property_id=property_id))
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                try:
                    import openpyxl
                    # Load the workbook
                    wb = openpyxl.load_workbook(file_path, data_only=True)
                    # Process each sheet
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        extracted_text += f"Sheet: {sheet_name}\n"
                        # Process each row
                        for row in ws.iter_rows():
                            row_text = ""
                            # Process each cell in the row
                            for cell in row:
                                if cell.value is not None:
                                    row_text += str(cell.value) + "\t"
                            if row_text:
                                extracted_text += row_text + "\n"
                        extracted_text += "\n\n"  # Separate sheets
                except Exception as xlsx_err:
                    print(f"Error extracting XLSX text: {xlsx_err}")
                    flash('Error processing Excel document.', 'danger')
                    os.remove(file_path)  # Clean up file
                    return redirect(url_for('views.knowledge_base', property_id=property_id))
            else:
                print(f"Unsupported MIME type: {mime_type}")
                flash('Unsupported file type.', 'danger')
                os.remove(file_path)  # Clean up file
                return redirect(url_for('views.knowledge_base', property_id=property_id))

            # Create knowledge source entry in Firestore
            import uuid
            source_id = str(uuid.uuid4())

            source_data = {
                'propertyId': property_id,
                'name': filename,
                'content': extracted_text,
                'type': 'file',
                'status': 'pending',
                'filePath': file_path,
                'sourceType': 'file',  # Add this field for the template display
                'fileName': filename,   # Add this field for the template display
                'createdAt': datetime.now(timezone.utc),
                'updatedAt': datetime.now(timezone.utc)
            }

            success = create_knowledge_source(source_id, source_data)

            if success:
                print(f"Created knowledge source {source_id} for property {property_id} in Firestore")

                # Auto-process the file to generate knowledge items
                try:
                    # Update source status to processing
                    update_status = {'status': 'processing', 'updatedAt': datetime.now(timezone.utc)}
                    update_knowledge_source(source_id, update_status)

                    # Generate knowledge items using Gemini with new schema
                    from concierge.api.utils import generate_knowledge_items_with_gemini
                    generated_items = generate_knowledge_items_with_gemini(extracted_text, property_details)

                    if not generated_items:
                        # No items generated
                        update_knowledge_source(source_id, {'status': 'processed_empty', 'updatedAt': datetime.now(timezone.utc)})
                        flash('File uploaded, but no knowledge items could be generated.', 'warning')
                        return redirect(url_for('views.knowledge_base', property_id=property_id))

                    # Import check_duplicate_content function
                    from concierge.utils.firestore_client import check_duplicate_content

                    # Create knowledge items in Firestore
                    item_count = 0
                    duplicate_count = 0

                    for item in generated_items:
                        if isinstance(item, dict) and item.get('content'):
                            # Check for duplicate content
                            content = item.get('content', '')
                            duplicate = check_duplicate_content(property_id, content)

                            if duplicate:
                                # Skip this item as it's a duplicate
                                print(f"Skipping duplicate content: {content[:50]}...")
                                duplicate_count += 1
                                continue

                            # Generate a unique ID for the item
                            item_id = str(uuid.uuid4())

                            # Create the item in Firestore with new schema
                            item_data = {
                                'propertyId': property_id,
                                'sourceId': source_id,
                                'type': item.get('type', 'information'),
                                'tags': item.get('tags', []),
                                'content': item.get('content', ''),
                                'status': 'pending',
                                'createdAt': datetime.now(timezone.utc),
                                'updatedAt': datetime.now(timezone.utc)
                            }

                            item_success = create_knowledge_item(item_id, item_data)
                            if item_success:
                                item_count += 1

                    if item_count > 0:
                        # Items created successfully
                        update_knowledge_source(source_id, {'status': 'processed', 'updatedAt': datetime.now(timezone.utc)})
                        if duplicate_count > 0:
                            flash(f'File uploaded: {item_count} knowledge items created, {duplicate_count} duplicates skipped.', 'success')
                        else:
                            flash(f'File uploaded and {item_count} knowledge items generated successfully. Please review them.', 'success')
                    else:
                        if duplicate_count > 0:
                            # All items were duplicates
                            update_knowledge_source(source_id, {'status': 'processed_duplicates', 'updatedAt': datetime.now(timezone.utc)})
                            flash(f'File uploaded, but all {duplicate_count} items were duplicates and skipped.', 'warning')
                        else:
                            # No valid items created
                            update_knowledge_source(source_id, {'status': 'processed_failed', 'updatedAt': datetime.now(timezone.utc)})
                            flash('File uploaded, but failed to create knowledge items.', 'danger')

                except Exception as process_err:
                    print(f"Error auto-processing file: {process_err}")
                    traceback.print_exc()

                    # Try to update source status
                    try:
                        update_knowledge_source(source_id, {'status': 'processing_error', 'updatedAt': datetime.now(timezone.utc)})
                    except:
                        pass

                    flash('File uploaded, but error processing knowledge items.', 'warning')
            else:
                print(f"Failed to create knowledge source for property {property_id} in Firestore")
                flash('Failed to save knowledge source to Firestore.', 'danger')
                os.remove(file_path)  # Clean up file

            return redirect(url_for('views.knowledge_base', property_id=property_id))
        else:
            flash('Unsupported file type.', 'warning')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

    except Exception as e:
        print(f"Error during knowledge upload for property {property_id} in Firestore: {e}")
        traceback.print_exc()
        flash(f'Error uploading file: {str(e)}', 'danger')
    return redirect(url_for('views.knowledge_base', property_id=property_id))

@views_bp.route('/properties/<property_id>/knowledge/add_text', methods=['POST'])
@login_required
def knowledge_add_text(property_id):
    """Add text knowledge directly for a property using Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_property, create_knowledge_source, update_knowledge_source, create_knowledge_item
    import traceback
    from datetime import datetime, timezone

    try:
        # --- Verify Ownership (Crucial for POST actions) ---
        property_data = get_property(property_id)

        # Check if user is the owner (handle both Firestore and DynamoDB formats)
        host_id = property_data.get('hostId', property_data.get('HostId')) if property_data else None
        if not property_data or host_id != user_id:
            print(f"Unauthorized text add attempt for property {property_id} by user {user_id}")
            flash('Permission denied or property not found.', 'danger')
            return redirect(url_for('views.properties_list'))
        # --- End Verification ---

        property_details = property_data # Already have the property details
        property_details['id'] = property_id # Add id for convenience

        # Check form data
        source_name = request.form.get('source_name')
        knowledge_text = request.form.get('knowledge_text')

        if not knowledge_text or not knowledge_text.strip():
            flash('Knowledge text is required.', 'warning')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Auto-generate source name if not provided
        if not source_name or not source_name.strip():
            # Generate source name from first 50 characters of content
            content_preview = knowledge_text.strip()[:50]
            # Remove newlines and extra spaces
            content_preview = ' '.join(content_preview.split())
            source_name = f"Text: {content_preview}..." if len(knowledge_text.strip()) > 50 else f"Text: {content_preview}"

        # Create knowledge source entry in Firestore
        import uuid
        source_id = str(uuid.uuid4())

        source_data = {
            'propertyId': property_id,
            'name': source_name,
            'content': knowledge_text,
            'type': 'text',
            'status': 'pending',
            'createdAt': datetime.now(timezone.utc),
            'updatedAt': datetime.now(timezone.utc)
        }

        success = create_knowledge_source(source_id, source_data)

        # If Firestore create fails, try DynamoDB as fallback
        if not success:
            from concierge.utils.dynamodb_client import create_knowledge_source as dynamo_create_source
            from concierge.utils.dynamodb_client import update_knowledge_source as dynamo_update_source
            from concierge.utils.dynamodb_client import create_knowledge_item as dynamo_create_item

            # Convert to DynamoDB format
            dynamo_source_data = {
                'PropertyId': property_id,
                'Name': source_name,
                'Content': knowledge_text,
                'Type': 'text',
                'Status': 'pending'
            }
            success = dynamo_create_source(source_id, dynamo_source_data)

        if success:
            print(f"Created text knowledge source {source_id} for property {property_id}")

            # --- AUTO-PROCESS: Generate Q&A pairs ---
            try:
                # Update source status to processing
                update_status = {'status': 'processing', 'updatedAt': datetime.now(timezone.utc)}
                update_success = update_knowledge_source(source_id, update_status)

                # If Firestore update fails, try DynamoDB
                if not update_success and 'dynamo_update_source' in locals():
                    dynamo_update_source(source_id, {'Status': 'processing'})

                # Generate knowledge items using Gemini with new schema
                from concierge.api.utils import generate_knowledge_items_with_gemini
                generated_items = generate_knowledge_items_with_gemini(knowledge_text, property_details)

                if not generated_items:
                    # No items generated
                    empty_status = {'status': 'processed_empty', 'updatedAt': datetime.now(timezone.utc)}
                    update_knowledge_source(source_id, empty_status)

                    # If Firestore update fails, try DynamoDB
                    if not update_success and 'dynamo_update_source' in locals():
                        dynamo_update_source(source_id, {'Status': 'processed_empty'})

                    flash('Knowledge text added, but no items could be generated.', 'warning')
                    return redirect(url_for('views.knowledge_base', property_id=property_id))

                # Import check_duplicate_content function
                from concierge.utils.firestore_client import check_duplicate_content

                # Create knowledge items in Firestore
                item_count = 0
                duplicate_count = 0

                for item in generated_items:
                    if isinstance(item, dict) and item.get('content'):
                        # Check for duplicate content
                        content = item.get('content', '')
                        duplicate = check_duplicate_content(property_id, content)

                        if duplicate:
                            # Skip this item as it's a duplicate
                            print(f"Skipping duplicate content: {content[:50]}...")
                            duplicate_count += 1
                            continue

                        # Generate a unique ID for the item
                        item_id = str(uuid.uuid4())

                        # Create the item in Firestore with new schema
                        item_data = {
                            'propertyId': property_id,
                            'sourceId': source_id,
                            'type': item.get('type', 'information'),
                            'tags': item.get('tags', []),
                            'content': item.get('content', ''),
                            'status': 'pending',
                            'createdAt': datetime.now(timezone.utc),
                            'updatedAt': datetime.now(timezone.utc)
                        }

                        item_success = create_knowledge_item(item_id, item_data)
                        if item_success:
                            item_count += 1

                if item_count > 0:
                    # Items created successfully
                    processed_status = {'status': 'processed', 'updatedAt': datetime.now(timezone.utc)}
                    update_knowledge_source(source_id, processed_status)

                    # If Firestore update fails, try DynamoDB
                    if not update_success and 'dynamo_update_source' in locals():
                        dynamo_update_source(source_id, {'Status': 'processed'})

                    if duplicate_count > 0:
                        flash(f'Text added: {item_count} knowledge items created, {duplicate_count} duplicates skipped.', 'success')
                    else:
                        flash(f'Text added and {item_count} knowledge items generated successfully. Please review them.', 'success')
                else:
                    if duplicate_count > 0:
                        # All items were duplicates
                        duplicate_status = {'status': 'processed_duplicates', 'updatedAt': datetime.now(timezone.utc)}
                        update_knowledge_source(source_id, duplicate_status)

                        # If Firestore update fails, try DynamoDB
                        if not update_success and 'dynamo_update_source' in locals():
                            dynamo_update_source(source_id, {'Status': 'processed_duplicates'})

                        flash(f'Text added, but all {duplicate_count} items were duplicates and skipped.', 'warning')
                    else:
                        # No valid items created
                        failed_status = {'status': 'processed_failed', 'updatedAt': datetime.now(timezone.utc)}
                        update_knowledge_source(source_id, failed_status)

                        # If Firestore update fails, try DynamoDB
                        if not update_success and 'dynamo_update_source' in locals():
                            dynamo_update_source(source_id, {'Status': 'processed_failed'})

                        flash('Text added, but failed to create knowledge items.', 'danger')
            except Exception as process_err:
                print(f"Error auto-processing knowledge text: {process_err}")
                traceback.print_exc()

                # Try to update source status
                try:
                    error_status = {'status': 'processing_error', 'updatedAt': datetime.now(timezone.utc)}
                    update_knowledge_source(source_id, error_status)

                    # If Firestore update fails, try DynamoDB
                    if 'dynamo_update_source' in locals():
                        dynamo_update_source(source_id, {'Status': 'processing_error'})
                except:
                    pass

                flash('Text added, but error processing Q&A pairs. Please try manual processing.', 'warning')
        else:
            print(f"Failed to create text knowledge source for property {property_id}")
            flash('Failed to save knowledge source to database.', 'danger')

        return redirect(url_for('views.knowledge_base', property_id=property_id))

    except Exception as e:
        print(f"Error during knowledge text addition for property {property_id}: {e}")
        traceback.print_exc()
        flash(f'Error adding text: {str(e)}', 'danger')
    return redirect(url_for('views.knowledge_base', property_id=property_id))

@views_bp.route('/properties/<property_id>/knowledge/process', methods=['POST'])
@login_required
def knowledge_process(property_id):
    """Process a knowledge source to generate Q&A pairs using DynamoDB"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.dynamodb_client import get_property, get_knowledge_source, update_knowledge_source, create_knowledge_item

    try:
        # Get the source ID from form
        source_id = request.form.get('source_id')
        if not source_id:
            flash('Source ID is required.', 'warning')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Verify property ownership
        property_data = get_property(property_id)
        if not property_data or property_data.get('hostId') != user_id:
            print(f"Unauthorized knowledge process attempt for property {property_id} by user {user_id}")
            flash('Permission denied or property not found.', 'danger')
            return redirect(url_for('views.properties_list'))

        # Get the knowledge source
        source_data = get_knowledge_source(source_id)
        if not source_data:
            flash('Knowledge source not found.', 'danger')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Verify source belongs to property
        source_property_id = source_data.get('PK', '').replace('PROPERTY#', '')
        if source_property_id != property_id:
            print(f"Source {source_id} does not belong to property {property_id}")
            flash('Knowledge source does not belong to this property.', 'danger')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Get the content to process
        content = source_data.get('Content', '')
        if not content:
            flash('No content to process.', 'warning')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Update source status to processing
        update_knowledge_source(source_id, {'Status': 'processing'})

        # Generate Q&A pairs using Gemini
        from concierge.api.utils import generate_qna_with_gemini
        generated_qna = generate_qna_with_gemini(content, property_data)

        if not generated_qna:
            # No Q&A pairs generated
            update_knowledge_source(source_id, {'Status': 'processed_empty'})
            flash('No Q&A pairs could be generated.', 'warning')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Create knowledge items in DynamoDB
        item_count = 0
        for qna in generated_qna:
            if isinstance(qna, dict) and qna.get('question') and qna.get('answer'):
                # Generate a unique ID for the item
                import uuid
                item_id = str(uuid.uuid4())

                # Create the item
                item_data = {
                    'propertyId': property_id,
                    'sourceId': source_id,
                    'question': qna['question'],
                    'answer': qna['answer'],
                    'status': 'pending',
                }

                success = create_knowledge_item(item_id, item_data)
                if success:
                    item_count += 1

        if item_count > 0:
            # Items created successfully
            update_knowledge_source(source_id, {'Status': 'processed'})
            flash(f'{item_count} Q&A pairs generated successfully.', 'success')
        else:
            # No valid items created
            update_knowledge_source(source_id, {'Status': 'processed_failed'})
            flash('Failed to create Q&A pairs.', 'danger')

        return redirect(url_for('views.knowledge_base', property_id=property_id))

    except Exception as e:
        print(f"Error processing knowledge source: {e}")
        traceback.print_exc()

        # Try to update source status if source_id is available
        if source_id:
            try:
                update_knowledge_source(source_id, {'Status': 'processing_error'})
            except:
                pass

        flash(f'Error processing knowledge source: {str(e)}', 'danger')
    return redirect(url_for('views.knowledge_base', property_id=property_id))

@views_bp.route('/properties/<property_id>/knowledge/review', methods=['GET'])
@login_required
def knowledge_review(property_id):
    """Redirect to unified knowledge management page"""
    # Redirect to the unified knowledge management page
    return redirect(url_for('views.knowledge_base', property_id=property_id))

@views_bp.route('/properties/<property_id>/knowledge/delete_source/<source_id>', methods=['POST'])
@login_required
def delete_knowledge_source(property_id, source_id):
    """Delete a knowledge source and all its items"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.dynamodb_client import get_property, get_knowledge_source, delete_knowledge_source

    try:
        # Verify property ownership
        property_data = get_property(property_id)
        if not property_data or property_data.get('hostId') != user_id:
            print(f"Unauthorized knowledge source deletion attempt for property {property_id} by user {user_id}")
            flash('Permission denied or property not found.', 'danger')
            return redirect(url_for('views.properties_list'))

        # Verify the source belongs to this property
        source_data = get_knowledge_source(source_id)
        if not source_data:
            flash('Knowledge source not found.', 'danger')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        source_property_id = source_data.get('PK', '').replace('PROPERTY#', '')
        if source_property_id != property_id:
            print(f"Source {source_id} does not belong to property {property_id}")
            flash('Knowledge source does not belong to this property.', 'danger')
            return redirect(url_for('views.knowledge_base', property_id=property_id))

        # Delete the source and its items
        success = delete_knowledge_source(source_id)

        if success:
            flash('Knowledge source deleted successfully.', 'success')
        else:
            flash('Failed to delete knowledge source.', 'danger')

        return redirect(url_for('views.knowledge_base', property_id=property_id))

    except Exception as e:
        print(f"Error deleting knowledge source {source_id} for property {property_id}: {e}")
        traceback.print_exc()
        flash(f'Error deleting knowledge source: {str(e)}', 'danger')
        return redirect(url_for('views.knowledge_base', property_id=property_id))

@views_bp.route('/properties/<property_id>/knowledge/approve_qna', methods=['POST'])
@login_required
def knowledge_approve_qna(property_id):
    """Approve or reject knowledge items using Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_property, get_knowledge_item, update_knowledge_item, update_knowledge_item_status

    try:
        # Verify property ownership
        property_data = get_property(property_id)
        if not property_data or property_data.get('hostId') != user_id:
            logger.warning(f"Unauthorized knowledge approval attempt for property {property_id} by user {user_id}")
            return jsonify({"success": False, "error": "Permission denied or property not found"})

        # Parse JSON data from request
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "Missing request data"})

        # Get required fields
        item_id = data.get('qna_id')
        action = data.get('status', '').lower()
        question = data.get('question')  # This will be used as a tag
        answer = data.get('answer')      # This will be used as content

        if not item_id:
            return jsonify({"success": False, "error": "Item ID is required"})

        # Get the knowledge item
        item_data = get_knowledge_item(item_id)
        if not item_data:
            logger.warning(f"Knowledge item {item_id} not found")
            return jsonify({"success": False, "error": "Knowledge item not found"})

        # Verify item belongs to property
        item_property_id = item_data.get('propertyId')
        if item_property_id != property_id:
            logger.warning(f"Item {item_id} does not belong to property {property_id}")
            return jsonify({"success": False, "error": "Knowledge item does not belong to this property"})

        # Prepare update data with edited tags and content
        # Convert question to tag and answer to content for new schema
        tags = [question] if question else []
        if item_data.get('tags'):
            # Keep existing tags if they exist, but replace the first one with the question
            tags = item_data.get('tags', [])
            if tags and question:
                tags[0] = question

        update_data = {
            'tags': tags,
            'content': answer
        }

        # Process the action
        if action == 'approved':
            # Also update status to approved
            update_data['status'] = 'approved'
            success = update_knowledge_item(item_id, update_data)
            if success:
                logger.info(f"Knowledge item {item_id} approved and updated by user {user_id}")
                return jsonify({"success": True, "message": "Knowledge item approved"})
            else:
                return jsonify({"success": False, "error": "Failed to update knowledge item"})
        elif action == 'rejected':
            # Update status to rejected
            update_data['status'] = 'rejected'
            success = update_knowledge_item(item_id, update_data)
            if success:
                logger.info(f"Knowledge item {item_id} rejected by user {user_id}")
                return jsonify({"success": True, "message": "Knowledge item rejected"})
            else:
                return jsonify({"success": False, "error": "Failed to update knowledge item"})
        else:
            return jsonify({"success": False, "error": "Invalid action"})

    except Exception as e:
        logger.error(f"Error processing knowledge approval: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Error processing request: {str(e)}"})

@views_bp.route('/api/knowledge/<string:item_id>/delete', methods=['DELETE'])
@login_required
def delete_knowledge_item(item_id):
    """Delete a knowledge item using Firestore"""
    user_id = session.get('user_id')

    # Import required functions
    from concierge.utils.dynamodb_client import get_knowledge_item, get_property, delete_knowledge_item as dynamo_delete_item

    if not user_id:
        return jsonify({"success": False, "error": "Authentication required"}), 401

    try:
        print(f"Processing deletion request for knowledge item {item_id} by user {user_id}")

        # Get the knowledge item
        item_data = get_knowledge_item(item_id)
        if not item_data:
            print(f"Knowledge item {item_id} not found for deletion")
            return jsonify({"success": False, "error": "Knowledge item not found"}), 404

        # Get the property ID from PK field (format: "PROPERTY#property_id")
        property_id = None
        pk = item_data.get('PK', '')
        if pk and pk.startswith('PROPERTY#'):
            property_id = pk.replace('PROPERTY#', '')

        if not property_id:
            print(f"Invalid knowledge item {item_id} (cannot extract property ID from {pk})")
            return jsonify({"success": False, "error": f"Invalid knowledge item (cannot extract property ID from {pk})"}), 400

        property_data = get_property(property_id)
        if not property_data:
            print(f"Property {property_id} not found for knowledge item {item_id}")
            return jsonify({"success": False, "error": "Property not found"}), 404

        # Check if the user is authorized to delete this item
        if property_data.get('HostId') != user_id:
            print(f"Unauthorized attempt to delete knowledge item {item_id} for property {property_id} by user {user_id}")
            return jsonify({"success": False, "error": "You don't have permission to delete this item"}), 403

        # Delete the item from DynamoDB
        success = dynamo_delete_item(item_id)
        if success:
            print(f"Knowledge item {item_id} for property {property_id} successfully deleted by user {user_id}")
            return jsonify({"success": True, "message": "Knowledge item deleted successfully"})
        else:
            print(f"Failed to delete knowledge item {item_id} for property {property_id}")
            return jsonify({"success": False, "error": "Failed to delete knowledge item"}), 500

    except Exception as e:
        print(f"Error deleting knowledge item {item_id}: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"An error occurred: {str(e)}"}), 500

@views_bp.route('/api/knowledge/<string:item_id>/update', methods=['PUT'])
@login_required
def update_knowledge_item(item_id):
    """Update a knowledge item using DynamoDB"""
    user_id = session.get('user_id')

    # Import required functions
    from concierge.utils.dynamodb_client import get_knowledge_item, get_property, update_knowledge_item as dynamo_update_item

    if not user_id:
        return jsonify({"success": False, "error": "Authentication required"}), 401

    # Get the JSON data
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400

    # Check for required fields
    question = data.get('question')
    answer = data.get('answer')

    if not question or not answer:
        return jsonify({"success": False, "error": "Question and answer are required"}), 400

    try:
        # Get the knowledge item
        item_data = get_knowledge_item(item_id)
        if not item_data:
            return jsonify({"success": False, "error": "Knowledge item not found"}), 404

        # Get the property ID from PK field (format: "PROPERTY#property_id")
        property_id = None
        pk = item_data.get('PK', '')
        if pk and pk.startswith('PROPERTY#'):
            property_id = pk.replace('PROPERTY#', '')

        if not property_id:
            return jsonify({"success": False, "error": f"Invalid knowledge item (cannot extract property ID from {pk})"}), 400

        property_data = get_property(property_id)
        if not property_data:
            return jsonify({"success": False, "error": "Property not found"}), 404

        # Check if the user is authorized to update this item
        if property_data.get('HostId') != user_id:
            print(f"Unauthorized attempt to update knowledge item {item_id} by user {user_id}")
            return jsonify({"success": False, "error": "You don't have permission to update this item"}), 403

        # Prepare the update data
        update_data = {
            'Question': question,
            'Answer': answer
        }

        # If status is provided and valid, update it
        status = data.get('status')
        valid_statuses = ['pending', 'approved', 'rejected']
        if status and status in valid_statuses:
            update_data['Status'] = status

        # Update the item
        success = dynamo_update_item(item_id, update_data)
        if success:
            print(f"Knowledge item {item_id} updated by user {user_id}")
            return jsonify({"success": True, "message": "Knowledge item updated successfully"})
        else:
            return jsonify({"success": False, "error": "Failed to update knowledge item"}), 500

    except Exception as e:
        print(f"Error updating knowledge item {item_id}: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"An error occurred: {str(e)}"}), 500

@views_bp.route('/properties/<string:property_id>/knowledge/manage')
@login_required
def knowledge_manage(property_id):
    """Redirect to unified knowledge management page"""
    # Redirect to the unified knowledge management page
    return redirect(url_for('views.knowledge_base', property_id=property_id))

# --- API Endpoint: Get Reservations for a Property (Host Only) ---
@views_bp.route('/properties/<property_id>/reservations', methods=['GET'])
@login_required
def property_reservations(property_id):
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_property, list_property_reservations

    # Fetch property and check ownership
    property_data = get_property(property_id)
    if not property_data:
        flash('Property not found in Firestore.', 'danger')
        return redirect(url_for('views.properties_list'))

    # Add ID to property data if not already present
    if 'id' not in property_data:
        property_data['id'] = property_id

    if property_data.get('hostId') != user_id:
        flash('Access denied: You do not own this property.', 'danger')
        return redirect(url_for('views.properties_list'))

    # Fetch reservations for this property from Firestore
    reservations = []
    try:
        all_reservations = list_property_reservations(property_id)

        # Import date utilities for consistent formatting
        from concierge.utils.date_utils import ensure_date_only_format

        for res_data in all_reservations:
            # Normalize dates first using the new utilities
            start_date_raw = (res_data.get('startDate') or res_data.get('StartDate') or 
                             res_data.get('checkInDate') or res_data.get('CheckInDate'))
            end_date_raw = (res_data.get('endDate') or res_data.get('EndDate') or 
                           res_data.get('checkOutDate') or res_data.get('CheckOutDate'))

            # Extract Airbnb URL from description if it exists
            airbnb_url = ''
            description = res_data.get('description', '')
            if description and 'Reservation URL:' in description:
                # Extract URL from description like "Reservation URL: https://www.airbnb.com/hosting/reservations/details/HMCKAYA8KN Phone Number..."
                import re
                url_match = re.search(r'Reservation URL:\s*(https://[^\s]+)', description)
                if url_match:
                    airbnb_url = url_match.group(1)

            # Format the response data with consistent date-only format
            formatted_res = {
                'id': res_data.get('id', ''),
                'startDate': ensure_date_only_format(start_date_raw) or '',
                'endDate': ensure_date_only_format(end_date_raw) or '',
                'guestName': res_data.get('guestName') or res_data.get('GuestName') or '',
                'guestPhoneNumber': res_data.get('guestPhoneNumber') or res_data.get('GuestPhoneNumber') or '',
                'guestPhoneLast4': res_data.get('guestPhoneLast4') or res_data.get('GuestPhoneLast4') or '',
                'summary': res_data.get('summary', ''),
                'description': res_data.get('description', ''),
                'fetchedAt': res_data.get('updatedAt', ''),
                'airbnb_url': airbnb_url,
                'AdditionalContacts': res_data.get('additional_contacts') or res_data.get('additionalContacts') or res_data.get('AdditionalContacts') or [],
            }
            reservations.append(formatted_res)

        # Sort reservations by start date (ascending, None last)
        reservations.sort(key=lambda r: r.get('startDate') or '9999-12-31')

    except Exception as e:
        print(f"Error fetching reservations from Firestore for property {property_id}: {e}")
        traceback.print_exc()
        flash('Failed to load reservations from Firestore.', 'danger')

    return render_template('property_reservations.html', property=property_data, reservations=reservations)

@views_bp.route('/api/properties/<property_id>/reservations', methods=['GET'])
@login_required
def get_property_reservations(property_id):
    user_id = session.get('user_id')
    user_role = session.get('user_role')
    if not user_id:
        return jsonify({'error': 'User not logged in'}), 401
    if user_role != 'host':
        return jsonify({'error': 'User is not a host'}), 403

    # Import required functions
    from concierge.utils.firestore_client import get_property, list_property_reservations

    # Check host owns property
    property_data = get_property(property_id)
    if not property_data:
        return jsonify({'error': 'Property not found in Firestore'}), 404

    if property_data.get('hostId') != user_id:
        return jsonify({'error': 'Access denied: not property host'}), 403

    # Fetch reservations for this property from Firestore
    reservations = []
    try:
        all_reservations = list_property_reservations(property_id)

        # Import date utilities for consistent formatting
        from concierge.utils.date_utils import ensure_date_only_format

        for res_data in all_reservations:
            # Normalize dates first using the new utilities
            start_date_raw = (res_data.get('startDate') or res_data.get('StartDate') or 
                             res_data.get('checkInDate') or res_data.get('CheckInDate'))
            end_date_raw = (res_data.get('endDate') or res_data.get('EndDate') or 
                           res_data.get('checkOutDate') or res_data.get('CheckOutDate'))

            # Extract Airbnb URL from description if it exists
            airbnb_url = ''
            description = res_data.get('description', '')
            if description and 'Reservation URL:' in description:
                # Extract URL from description like "Reservation URL: https://www.airbnb.com/hosting/reservations/details/HMCKAYA8KN Phone Number..."
                import re
                url_match = re.search(r'Reservation URL:\s*(https://[^\s]+)', description)
                if url_match:
                    airbnb_url = url_match.group(1)

            # Format the response data with consistent date-only format
            formatted_res = {
                'id': res_data.get('id', ''),
                'startDate': ensure_date_only_format(start_date_raw) or '',
                'endDate': ensure_date_only_format(end_date_raw) or '',
                'guestName': res_data.get('guestName') or res_data.get('GuestName') or '',
                'guestPhoneNumber': res_data.get('guestPhoneNumber') or res_data.get('GuestPhoneNumber') or '',
                'guestPhoneLast4': res_data.get('guestPhoneLast4') or res_data.get('GuestPhoneLast4') or '',
                'summary': res_data.get('summary', ''),
                'description': res_data.get('description', ''),
                'fetchedAt': res_data.get('updatedAt', ''),
                'airbnb_url': airbnb_url,
                'AdditionalContacts': res_data.get('additional_contacts') or res_data.get('additionalContacts') or res_data.get('AdditionalContacts') or [],
            }
            reservations.append(formatted_res)

        return jsonify({'reservations': reservations})
    except Exception as e:
        print(f"Error fetching reservations from Firestore for property {property_id}: {e}")
        import traceback; traceback.print_exc()
        return jsonify({'error': 'Failed to fetch reservations from Firestore'}), 500

# --- API Endpoint to Update Reservation Phone --- #
@views_bp.route('/api/property/<property_id>/reservations/<reservation_id>/phone', methods=['PUT'])
@login_required
def update_reservation_phone(property_id, reservation_id):
    """Update the phone number for a reservation (Legacy endpoint for compatibility)."""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.dynamodb_client import get_property, update_reservation_phone as dynamo_update_phone

    try:
        # Get the phone number from request
        data = request.json
        if not data or 'phone_number' not in data:
            return jsonify({'success': False, 'error': 'No phone number provided'}), 400

        phone_number = data['phone_number']

        # Verify property ownership
        property_data = get_property(property_id)
        if not property_data:
            return jsonify({'success': False, 'error': 'Property not found'}), 404

        if property_data.get('HostId') != user_id:
            return jsonify({'success': False, 'error': 'Not authorized to update this property'}), 403

        # Update the phone number
        success = dynamo_update_phone(reservation_id, phone_number)

        if success:
            return jsonify({
                'success': True,
                'message': 'Phone number updated successfully'
            })
        else:
            return jsonify({'success': False, 'error': 'Failed to update phone number'}), 500

    except Exception as e:
        print(f"Error updating phone for reservation {reservation_id}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@views_bp.route('/api/property/<property_id>/reservations/<reservation_id>/contacts', methods=['PUT'])
@login_required
def update_reservation_contacts(property_id, reservation_id):
    """Update multiple contact information for a reservation."""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_property, update_reservation_contacts

    # Verify user has access to this property
    try:
        property_data = get_property(property_id)

        if not property_data:
            return jsonify({'success': False, 'error': 'Property not found'}), 404

        if property_data.get('hostId') != user_id:
            return jsonify({'success': False, 'error': 'Not authorized to update this property'}), 403

        # Get the contacts from request
        data = request.json
        if not data or 'contacts' not in data:
            return jsonify({'success': False, 'error': 'No contacts provided'}), 400

        contacts = data['contacts']
        # Validate contacts format
        if not isinstance(contacts, list):
            return jsonify({'success': False, 'error': 'Contacts must be an array'}), 400

        for contact in contacts:
            if not isinstance(contact, dict) or 'name' not in contact or 'phone' not in contact:
                return jsonify({'success': False, 'error': 'Each contact must have name and phone'}), 400

        # Use the Firestore function to update contacts
        success = update_reservation_contacts(reservation_id, contacts)

        if success:
            return jsonify({
                'success': True,
                'message': 'Contacts updated successfully'
            })
        else:
            return jsonify({'success': False, 'error': 'Failed to update contacts'}), 500

    except Exception as e:
        print(f"Error updating contacts for reservation {reservation_id}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@views_bp.route('/api/property/<property_id>/reservations/refresh', methods=['POST'])
@login_required
def refresh_property_reservations(property_id):
    """Manually trigger a refresh of reservations for a specific property from iCal."""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_property, list_property_reservations
    from concierge.utils.reservations import fetch_and_parse_ical, update_all_reservations

    try:
        # Verify property ownership
        property_data = get_property(property_id)
        if not property_data:
            return jsonify({'success': False, 'error': 'Property not found'}), 404

        if property_data.get('hostId') != user_id:
            return jsonify({'success': False, 'error': 'Not authorized to refresh this property'}), 403

        # Check if property has an iCal URL
        ical_url = property_data.get('icalUrl')
        if not ical_url:
            return jsonify({'success': False, 'error': 'Property does not have an iCal URL configured'}), 400

        # Trigger the reservation update for this specific property
        # We'll create a simplified version of update_all_reservations for a single property
        from concierge.utils.firestore_client import (
            get_firestore_client,
            create_reservation,
            update_reservation,
            delete_reservation
        )
        from datetime import datetime, timezone

        print(f"[Manual Refresh] Processing property {property_id} with URL: {ical_url}")

        # Fetch events from iCal
        fetched_events = fetch_and_parse_ical(ical_url)

        if fetched_events is None:
            return jsonify({'success': False, 'error': 'Failed to fetch or parse iCal data'}), 500

        # Get existing reservations for this property
        existing_reservations = list_property_reservations(property_id)
        existing_by_date = {}

        # Group existing reservations by dates for easier matching
        for reservation in existing_reservations:
            start_date = reservation.get('startDate')
            end_date = reservation.get('endDate')

            if start_date and end_date:
                date_key = f"{start_date}_{end_date}"
                reservation_id = reservation.get('id')
                existing_by_date[date_key] = {
                    'id': reservation_id,
                    'data': reservation
                }

        # Process fetched events
        processed_ids = set()
        updates_count = 0
        added_count = 0

        for event in fetched_events:
            start_date = event.get('start')
            end_date = event.get('end')

            if not (start_date and end_date):
                continue

            date_key = f"{start_date}_{end_date}"

            # Check if this reservation already exists
            if date_key in existing_by_date:
                # Update existing reservation
                existing_item = existing_by_date[date_key]
                existing_data = existing_item['data']
                reservation_id = existing_item['id']

                processed_ids.add(reservation_id)

                # Preserve existing contact information
                guest_phone_number = existing_data.get('guestPhoneNumber')
                additional_contacts = (
                    existing_data.get('additional_contacts') or
                    existing_data.get('additionalContacts') or
                    existing_data.get('AdditionalContacts') or
                    []
                )
                guest_name = existing_data.get('guestName')

                # Check if any fields need updating
                update_needed = False
                update_data = {}

                if existing_data.get('summary') != event.get('summary'):
                    update_data['summary'] = event.get('summary')
                    update_needed = True

                if existing_data.get('description') != event.get('description'):
                    update_data['description'] = event.get('description')
                    update_needed = True

                if not guest_phone_number and event.get('phone_last_4') and existing_data.get('guestPhoneLast4') != event.get('phone_last_4'):
                    update_data['guestPhoneLast4'] = event.get('phone_last_4')
                    update_needed = True

                # Preserve additional contacts if they exist
                if additional_contacts:
                    update_data['additional_contacts'] = additional_contacts
                    update_needed = True

                # Preserve guest name if it exists
                if guest_name:
                    update_data['guestName'] = guest_name

                # Preserve guest phone number if it exists
                if guest_phone_number:
                    update_data['guestPhoneNumber'] = guest_phone_number

                # CRITICAL FIX: Always preserve existing dates to prevent them from being set to null
                # This prevents the date normalization process from nullifying dates when only
                # preserving additional contacts or other fields
                update_data['startDate'] = existing_data.get('startDate')
                update_data['endDate'] = existing_data.get('endDate')

                # Add update timestamp
                update_data['updatedAt'] = datetime.now(timezone.utc)

                if update_needed:
                    update_reservation(reservation_id, update_data)
                    updates_count += 1
            else:
                # This is a new reservation
                reservation_data = {
                    'propertyId': property_id,
                    'startDate': start_date,
                    'endDate': end_date,
                    'summary': event.get('summary'),
                    'description': event.get('description'),
                    'guestPhoneLast4': event.get('phone_last_4'),
                    'status': 'active',
                    'createdAt': datetime.now(timezone.utc),
                    'updatedAt': datetime.now(timezone.utc)
                }

                new_id = create_reservation(reservation_data)
                if new_id:
                    added_count += 1

        # Handle reservations to delete (those not in the fetched events)
        deleted_count = 0
        for reservation in existing_reservations:
            reservation_id = reservation.get('id')
            if reservation_id and reservation_id not in processed_ids:
                # Check if this reservation has custom contact info that should be preserved
                additional_contacts_to_check = (
                    reservation.get('additional_contacts') or
                    reservation.get('additionalContacts') or
                    reservation.get('AdditionalContacts') or
                    []
                )
                has_custom_contacts = (
                    reservation.get('guestPhoneNumber') or
                    (additional_contacts_to_check and len(additional_contacts_to_check) > 0)
                )

                if not has_custom_contacts:
                    delete_reservation(reservation_id)
                    deleted_count += 1

        return jsonify({
            'success': True,
            'message': f'Reservations refreshed successfully: {updates_count} updated, {added_count} added, {deleted_count} deleted',
            'stats': {
                'updated': updates_count,
                'added': added_count,
                'deleted': deleted_count,
                'total_events': len(fetched_events)
            }
        })

    except Exception as e:
        print(f"Error refreshing reservations for property {property_id}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


@views_bp.route('/api/active-reservations', methods=['GET'])
@login_required
def get_active_reservations():
    """Get all active reservations for the host's properties from Firestore."""
    user_id = getattr(g, 'user_id', None)

    try:
        # Get properties from Firestore
        from concierge.utils.firestore_client import list_properties_by_host, list_property_reservations
        host_properties = list_properties_by_host(user_id)

        # Extract property IDs from Firestore properties
        property_ids = [prop.get('id') for prop in host_properties if prop.get('id')]

        if not property_ids:
            print(f"No properties found in Firestore for user {user_id}")
            return jsonify({'success': True, 'reservations': []})

        # Get active reservations for all properties
        active_reservations = []

        # Get reservations from Firestore
        for property_id in property_ids:
            # Get reservations for this property from Firestore
            property_reservations = list_property_reservations(property_id)

            # Find corresponding property for name
            property_name = 'Unknown Property'
            for prop in host_properties:
                if prop.get('id') == property_id:
                    property_name = prop.get('name', 'Unknown Property')
                    break

            # Process each reservation
            for res in property_reservations:
                # Normalize dates first
                normalized_res = normalize_reservation_dates(res)

                # Only include active reservations (current or future)
                # TODO: Add date filtering logic here

                # Format the reservation data
                res_data = {
                    'id': res.get('id', ''),
                    'propertyId': property_id,
                    'propertyName': property_name,
                    'startDate': normalized_res.get('startDate', ''),
                    'endDate': normalized_res.get('endDate', ''),
                    'guestName': res.get('guestName', ''),
                    'guestPhoneNumber': res.get('guestPhoneNumber', '')
                }
                active_reservations.append(res_data)

        return jsonify({
            'success': True,
            'reservations': active_reservations
        })

    except Exception as e:
        print(f"Error getting active reservations from Firestore: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@views_bp.route('/properties_debug', methods=['GET'])
@login_required
def properties_debug():
    """Debug route to directly view properties data as JSON"""
    user_id = getattr(g, 'user_id', None)

    result = {
        'user_id': user_id,
        'raw_properties': [],
        'processed_properties': []
    }

    try:
        # First try to get properties from Firestore
        from concierge.utils.firestore_client import list_properties_by_host
        raw_properties = list_properties_by_host(user_id)

        # Ensure serializable for JSON
        for prop in raw_properties:
            serializable_prop = {}
            for k, v in prop.items():
                # Convert any non-serializable values to strings
                serializable_prop[k] = str(v)
            result['raw_properties'].append(serializable_prop)

        # Process properties as we do for the template
        processed_properties = []
        for prop in raw_properties:
            # Firestore already includes 'id' field
            property_id = prop.get('id', '')
            processed_prop = {
                'id': property_id,
                'name': prop.get('name', 'Unnamed Property'),
                'description': prop.get('description', 'No description'),
                'address': prop.get('address', 'N/A')
            }
            processed_properties.append(processed_prop)

        # All properties are now stored in Firestore only

        result['processed_properties'] = processed_properties
        result['source'] = 'firestore' if processed_properties else 'none'

    except Exception as e:
        result['error'] = str(e)
        import traceback
        traceback.print_exc()

    # Return as JSON
    return jsonify(result)

# API routes for knowledge items
@views_bp.route('/api/knowledge-items/<string:item_id>/delete', methods=['DELETE'])
@login_required
def delete_knowledge_item_firestore(item_id):
    """Delete a knowledge item using Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_knowledge_item, get_property, delete_knowledge_item

    if not user_id:
        return jsonify({"success": False, "error": "Authentication required"}), 401

    try:
        logger.info(f"Processing deletion request for knowledge item {item_id} by user {user_id}")

        # Get the knowledge item
        item_data = get_knowledge_item(item_id)
        if not item_data:
            logger.warning(f"Knowledge item {item_id} not found for deletion")
            return jsonify({"success": False, "error": "Knowledge item not found"}), 404

        # Get property ID from the item
        property_id = item_data.get('propertyId')
        if not property_id:
            logger.error(f"Invalid knowledge item {item_id} (cannot extract property ID from item)")
            return jsonify({"success": False, "error": "Invalid knowledge item (cannot extract property ID from item)"}), 400

        # Get property data to check ownership
        property_data = get_property(property_id)
        if not property_data:
            logger.error(f"Property {property_id} not found for knowledge item {item_id}")
            return jsonify({"success": False, "error": "Property not found"}), 404

        # Check if the user is authorized to delete this item
        if property_data.get('hostId') != user_id:
            logger.warning(f"Unauthorized attempt to delete knowledge item {item_id} by user {user_id}")
            return jsonify({"success": False, "error": "You don't have permission to delete this item"}), 403

        # Delete the item from DynamoDB
        success = delete_knowledge_item(item_id)
        if success:
            logger.info(f"Knowledge item {item_id} deleted by user {user_id}")
            return jsonify({"success": True, "message": "Knowledge item deleted successfully"})
        else:
            logger.error(f"Failed to delete knowledge item {item_id}")
            return jsonify({"success": False, "error": "Failed to delete knowledge item"}), 500

    except Exception as e:
        logger.error(f"Error deleting knowledge item {item_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": f"An error occurred: {str(e)}"}), 500

@views_bp.route('/api/knowledge-items/<string:item_id>/update', methods=['PUT'])
@login_required
def update_knowledge_item_firestore(item_id):
    """Update a knowledge item using Firestore"""
    user_id = getattr(g, 'user_id', None)

    # Import required functions
    from concierge.utils.firestore_client import get_knowledge_item, get_property, update_knowledge_item, generate_embedding

    if not user_id:
        return jsonify({"success": False, "error": "Authentication required"}), 401

    # Get the JSON data
    data = request.get_json()
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400

    try:
        # Extract fields from request
        item_type = data.get('type')
        tags = data.get('tags', [])
        content = data.get('content')

        # Validate required fields
        if not content:
            return jsonify({"success": False, "error": "Content is required"}), 400

        if not item_type:
            return jsonify({"success": False, "error": "Type is required"}), 400

        # Get the knowledge item
        item_data = get_knowledge_item(item_id)
        if not item_data:
            logger.warning(f"Knowledge item {item_id} not found for update")
            return jsonify({"success": False, "error": "Knowledge item not found"}), 404

        # Get property ID from the item
        property_id = item_data.get('propertyId')
        if not property_id:
            logger.error(f"Invalid knowledge item {item_id} (cannot extract property ID from item)")
            return jsonify({"success": False, "error": "Invalid knowledge item (cannot extract property ID from item)"}), 400

        # Get property data to check ownership
        property_data = get_property(property_id)
        if not property_data:
            logger.error(f"Property {property_id} not found for knowledge item {item_id}")
            return jsonify({"success": False, "error": "Property not found"}), 404

        # Check if the user is authorized to update this item
        if property_data.get('hostId') != user_id:
            logger.warning(f"Unauthorized attempt to update knowledge item {item_id} by user {user_id}")
            return jsonify({"success": False, "error": "You don't have permission to update this item"}), 403

        # Prepare the update data
        update_data = {
            'type': item_type,
            'tags': tags,
            'content': content
        }

        # If status is provided and valid, update it
        status = data.get('status')
        valid_statuses = ['pending', 'approved', 'rejected']
        if status and status in valid_statuses:
            update_data['status'] = status

        # Update the item
        success = update_knowledge_item(item_id, update_data)
        if success:
            logger.info(f"Knowledge item {item_id} updated by user {user_id}")
            return jsonify({"success": True, "message": "Knowledge item updated successfully"})
        else:
            logger.error(f"Failed to update knowledge item {item_id}")
            return jsonify({"success": False, "error": "Failed to update knowledge item"}), 500

    except Exception as e:
        logger.error(f"Error updating knowledge item {item_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": f"An error occurred: {str(e)}"}), 500

# API route for fetching knowledge items
@views_bp.route('/api/knowledge-items', methods=['GET'])
def get_property_knowledge_items():
    """Get knowledge items for a property from Firestore."""
    try:
        property_id = request.args.get('propertyId')
        status = request.args.get('status')  # Optional status filter

        if not property_id:
            return jsonify({'success': False, 'error': 'Property ID is required'}), 400

        # Get knowledge items from Firestore
        from concierge.utils.firestore_client import list_knowledge_items_by_property

        # If status is provided, filter by status
        if status:
            items = list_knowledge_items_by_property(property_id, status)
        else:
            items = list_knowledge_items_by_property(property_id)

        # Remove embedding field from items before returning as JSON
        # because Vector objects are not JSON serializable
        for item in items:
            if 'embedding' in item:
                del item['embedding']

        logger.info(f"Returning {len(items)} knowledge items for property {property_id}")

        return jsonify({
            'success': True,
            'items': items
        })

    except Exception as e:
        logger.error(f"Error fetching knowledge items from Firestore: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'Failed to fetch knowledge items: {str(e)}'
        }), 500

# --- Property Setup Wizard ---
@views_bp.route('/setup', methods=['GET', 'POST'])
@login_required
def property_setup_wizard():
    """Multi-step property setup wizard"""
    user_id = getattr(g, 'user_id', None)
    
    if request.method == 'POST':
        # Handle file uploads (multipart/form-data) differently from JSON data
        if request.content_type and 'multipart/form-data' in request.content_type:
            # Handle file upload
            action = request.form.get('action')
            
            if action == 'upload_file':
                try:
                    if 'file' not in request.files:
                        return jsonify({'success': False, 'error': 'No file provided'}), 400
                    
                    file = request.files['file']
                    if file.filename == '':
                        return jsonify({'success': False, 'error': 'No file selected'}), 400
                    
                    upload_type = request.form.get('upload_type', 'unknown')  # 'json' or 'facts'
                    
                    # Process the uploaded file
                    processed_data = process_uploaded_file(file, upload_type)
                    
                    return jsonify({
                        'success': True,
                        'data': processed_data,
                        'filename': file.filename
                    })
                    
                except Exception as e:
                    logger.error(f"Error processing file upload: {str(e)}")
                    return jsonify({'success': False, 'error': f'Failed to process file: {str(e)}'}), 500
            else:
                return jsonify({'success': False, 'error': 'Invalid file upload action'}), 400
        
        # Handle AJAX requests for saving wizard progress and data (JSON)
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'No data provided'}), 400
            
        action = data.get('action')
        
        if action == 'save_progress':
            # Save wizard progress to Firestore instead of session to avoid cookie size limits
            try:
                step = data.get('step', 1)
                wizard_data = data.get('data', {})
                
                # Get Firestore client
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                
                # Store in Firestore wizard_progress collection
                progress_doc = {
                    'user_id': user_id,
                    'step': step,
                    'data': wizard_data,
                    'last_saved': datetime.now(timezone.utc),
                    'updated_at': datetime.now(timezone.utc)
                }
                
                # Save to Firestore with user_id as document ID
                db.collection('wizard_progress').document(user_id).set(progress_doc)
                
                logger.info(f"🔵 WIZARD SAVE: User {user_id} at step {step} - Data size: {len(str(wizard_data))} chars")
                return jsonify({'success': True})
                
            except Exception as e:
                logger.error(f"Error saving wizard progress to Firestore: {str(e)}")
                return jsonify({'success': False, 'error': 'Failed to save progress'}), 500
            
        elif action == 'load_progress':
            # Load wizard progress from Firestore
            try:
                # Get Firestore client
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                
                # Load from Firestore wizard_progress collection
                progress_doc = db.collection('wizard_progress').document(user_id).get()
                
                if progress_doc.exists:
                    progress_data = progress_doc.to_dict()
                    logger.info(f"🔵 WIZARD LOAD: User {user_id} - Found data at step {progress_data.get('step', 1)}")
                    return jsonify({
                        'success': True, 
                        'data': progress_data.get('data', {}),
                        'step': progress_data.get('step', 1),
                        'last_saved': progress_data.get('last_saved').isoformat() if progress_data.get('last_saved') else None
                    })
                else:
                    logger.info(f"🔵 WIZARD LOAD: User {user_id} - No progress found")
                    return jsonify({'success': True, 'data': {}, 'step': 1})
                    
            except Exception as e:
                logger.error(f"Error loading wizard progress from Firestore: {str(e)}")
                return jsonify({'success': False, 'error': 'Failed to load progress'}), 500
            
        elif action == 'search_local_places':
            # Implement Google Places API search for local recommendations
            property_address = data.get('address', '')
            
            if not property_address:
                return jsonify({'success': False, 'error': 'Property address is required'}), 400
                
            try:
                places = search_local_places_google(property_address)
                return jsonify({
                    'success': True, 
                    'places': places
                })
            except Exception as e:
                logger.error(f"Error searching local places: {str(e)}")
                # Fallback to sample data if Google API fails
                return jsonify({
                    'success': True, 
                    'places': [
                        {'name': 'Local Restaurant', 'type': 'restaurant', 'selected': False},
                        {'name': 'Grocery Store', 'type': 'shopping', 'selected': False},
                        {'name': 'Coffee Shop', 'type': 'cafe', 'selected': False}
                    ]
                })
            

        elif action == 'create_property':
            # Create property from wizard data
            wizard_data = data.get('wizard_data', {})
            
            try:
                # Extract basic information
                basic_info = wizard_data.get('basic_info', {})
                property_name = basic_info.get('property_name', '').strip()
                property_address = basic_info.get('property_address', '').strip()
                
                if not property_name or not property_address:
                    return jsonify({'success': False, 'error': 'Property name and address are required'}), 400
                
                # Extract WiFi info from basic_info or voice sessions
                wifi_network = basic_info.get('wifi_network', '').strip()
                wifi_password = basic_info.get('wifi_password', '').strip()
                
                # Check voice sessions for WiFi info if not in basic_info
                if not wifi_network or not wifi_password:
                    voice_sessions = wizard_data.get('voice_sessions', [])
                    for session in voice_sessions:
                        voice_data = session.get('voiceCollectedData', [])
                        for item in voice_data:
                            if item.get('topic') == 'wifi':
                                if not wifi_network and item.get('networkName'):
                                    wifi_network = item['networkName']
                                if not wifi_password and item.get('password'):
                                    wifi_password = item['password']
                
                # Create property data structure
                property_data = {
                    'name': property_name,
                    'address': property_address,
                    'description': basic_info.get('description', '').strip(),
                    'checkInTime': basic_info.get('checkin_time', '15:00'),
                    'checkOutTime': basic_info.get('checkout_time', '11:00'),
                    'hostId': user_id,
                    'createdAt': datetime.now(timezone.utc),
                    'updatedAt': datetime.now(timezone.utc),
                    'status': 'active',
                    'setup_wizard_completed': True,
                    'wizard_data': wizard_data  # Store full wizard data for reference
                }
                
                # Add WiFi details if available
                if wifi_network or wifi_password:
                    property_data['wifiDetails'] = {
                        'network': wifi_network,
                        'password': wifi_password
                    }
                
                # Add Airbnb iCal URL if available
                airbnb_data = wizard_data.get('airbnb_data', {})
                if airbnb_data.get('ical_url'):
                    property_data['icalUrl'] = airbnb_data['ical_url']
                
                # Generate a unique property ID
                import uuid
                property_id = str(uuid.uuid4())
                
                # Create the property
                success = create_property(property_id, property_data)
                if not success:
                    return jsonify({'success': False, 'error': 'Failed to create property'}), 500
                
                # Process approved knowledge drafts into actual knowledge items
                knowledge_items_created = process_approved_knowledge_drafts(property_id, wizard_data)
                
                # Clean up wizard progress from Firestore after successful property creation
                try:
                    from concierge.utils.firestore_client import get_firestore_client
                    db = get_firestore_client()
                    db.collection('wizard_progress').document(user_id).delete()
                    logger.info(f"Wizard progress cleaned up for user {user_id} after property creation")
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup wizard progress for user {user_id}: {cleanup_error}")
                
                logger.info(f"Property {property_id} created successfully with {knowledge_items_created} knowledge items")
                return jsonify({
                    'success': True, 
                    'property_id': property_id,
                    'knowledge_items_created': knowledge_items_created
                })
                
            except Exception as e:
                logger.error(f"Error creating property from wizard data: {str(e)}")
                return jsonify({'success': False, 'error': f'Failed to create property: {str(e)}'}), 500

        elif action == 'generate_knowledge_drafts':
            # Generate draft knowledge items from wizard data using AI
            try:
                # Use wizard data from request if provided, otherwise load from Firestore
                wizard_data = data.get('wizard_data')
                
                if not wizard_data:
                    # Fallback: Load current wizard data from Firestore
                    from concierge.utils.firestore_client import get_firestore_client
                    db = get_firestore_client()
                    doc_ref = db.collection('wizard_progress').document(user_id)
                    doc = doc_ref.get()
                    
                    if not doc.exists:
                        return jsonify({'success': False, 'error': 'No wizard progress found'}), 404
                    
                    wizard_data = doc.to_dict()
                
                logger.info(f"Generating knowledge drafts for user {user_id} with {len(wizard_data)} sections")
                
                # Generate knowledge drafts using AI
                knowledge_drafts = generate_wizard_knowledge_drafts(wizard_data)
                
                if not knowledge_drafts:
                    return jsonify({'success': False, 'error': 'No knowledge items could be generated'}), 400
                
                # Store drafts back in the wizard progress document
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                doc_ref = db.collection('wizard_progress').document(user_id)
                
                # Load existing data to preserve other fields
                doc = doc_ref.get()
                if doc.exists:
                    existing_data = doc.to_dict()
                else:
                    existing_data = {}
                
                # Update with new drafts
                existing_data.update(wizard_data)
                existing_data['knowledge_drafts'] = knowledge_drafts
                existing_data['knowledge_drafts_generated_at'] = datetime.now(timezone.utc).isoformat()
                
                # Save updated wizard data
                doc_ref.set(existing_data)
                
                logger.info(f"Generated {len(knowledge_drafts)} knowledge drafts for user {user_id}")
                return jsonify({
                    'success': True,
                    'drafts_count': len(knowledge_drafts),
                    'drafts': knowledge_drafts
                })
                
            except Exception as e:
                logger.error(f"Error generating knowledge drafts for user {user_id}: {str(e)}")
                return jsonify({'success': False, 'error': f'Failed to generate knowledge drafts: {str(e)}'}), 500

        elif action == 'update_knowledge_draft':
            # Update a specific knowledge draft item
            try:
                draft_id = data.get('draft_id')
                updated_draft = data.get('draft_data')
                
                if not draft_id or not updated_draft:
                    return jsonify({'success': False, 'error': 'Missing required data'}), 400
                
                # Load current wizard data
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                doc_ref = db.collection('wizard_progress').document(user_id)
                doc = doc_ref.get()
                
                if not doc.exists:
                    return jsonify({'success': False, 'error': 'No wizard progress found'}), 404
                
                wizard_data = doc.to_dict()
                knowledge_drafts = wizard_data.get('knowledge_drafts', [])
                
                # Find and update the draft
                draft_updated = False
                for i, draft in enumerate(knowledge_drafts):
                    if draft.get('id') == draft_id:
                        knowledge_drafts[i] = {**draft, **updated_draft}
                        draft_updated = True
                        break
                
                if not draft_updated:
                    return jsonify({'success': False, 'error': 'Draft not found'}), 404
                
                # Save updated wizard data
                wizard_data['knowledge_drafts'] = knowledge_drafts
                wizard_data['knowledge_drafts_updated_at'] = datetime.now(timezone.utc).isoformat()
                doc_ref.set(wizard_data)
                
                return jsonify({'success': True})
                
            except Exception as e:
                logger.error(f"Error updating knowledge draft: {str(e)}")
                return jsonify({'success': False, 'error': str(e)}), 500

        elif action == 'delete_knowledge_draft':
            # Delete a specific knowledge draft item
            try:
                draft_id = data.get('draft_id')
                
                if not draft_id:
                    return jsonify({'success': False, 'error': 'Missing draft ID'}), 400
                
                # Load current wizard data
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                doc_ref = db.collection('wizard_progress').document(user_id)
                doc = doc_ref.get()
                
                if not doc.exists:
                    return jsonify({'success': False, 'error': 'No wizard progress found'}), 404
                
                wizard_data = doc.to_dict()
                knowledge_drafts = wizard_data.get('knowledge_drafts', [])
                
                # Remove the draft
                knowledge_drafts = [draft for draft in knowledge_drafts if draft.get('id') != draft_id]
                
                # Save updated wizard data
                wizard_data['knowledge_drafts'] = knowledge_drafts
                wizard_data['knowledge_drafts_updated_at'] = datetime.now(timezone.utc).isoformat()
                doc_ref.set(wizard_data)
                
                return jsonify({'success': True})
                
            except Exception as e:
                logger.error(f"Error deleting knowledge draft: {str(e)}")
                return jsonify({'success': False, 'error': str(e)}), 500

        elif action == 'add_knowledge_draft':
            # Add a new knowledge draft item
            try:
                draft_data = data.get('draft_data')
                
                if not draft_data:
                    return jsonify({'success': False, 'error': 'Missing draft data'}), 400
                
                # Load current wizard data
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                doc_ref = db.collection('wizard_progress').document(user_id)
                doc = doc_ref.get()
                
                if not doc.exists:
                    return jsonify({'success': False, 'error': 'No wizard progress found'}), 404
                
                wizard_data = doc.to_dict()
                knowledge_drafts = wizard_data.get('knowledge_drafts', [])
                
                # Add new draft with unique ID
                import uuid
                new_draft = {
                    'id': str(uuid.uuid4()),
                    'status': 'pending',
                    'created_at': datetime.now(timezone.utc).isoformat(),
                    **draft_data
                }
                
                knowledge_drafts.append(new_draft)
                
                # Save updated wizard data
                wizard_data['knowledge_drafts'] = knowledge_drafts
                wizard_data['knowledge_drafts_updated_at'] = datetime.now(timezone.utc).isoformat()
                doc_ref.set(wizard_data)
                
                return jsonify({'success': True, 'draft': new_draft})
                
            except Exception as e:
                logger.error(f"Error adding knowledge draft: {str(e)}")
                return jsonify({'success': False, 'error': str(e)}), 500

        elif action == 'save_knowledge_drafts':
            # Save knowledge drafts to wizard progress
            try:
                drafts = data.get('drafts')
                
                if not drafts:
                    return jsonify({'success': False, 'error': 'Missing drafts data'}), 400
                
                # Load current wizard data
                from concierge.utils.firestore_client import get_firestore_client
                db = get_firestore_client()
                doc_ref = db.collection('wizard_progress').document(user_id)
                doc = doc_ref.get()
                
                if not doc.exists:
                    return jsonify({'success': False, 'error': 'No wizard progress found'}), 404
                
                wizard_data = doc.to_dict()
                
                # Update knowledge drafts
                wizard_data['knowledge_drafts'] = drafts
                wizard_data['knowledge_drafts_updated_at'] = datetime.now(timezone.utc).isoformat()
                
                # Save updated wizard data
                doc_ref.set(wizard_data)
                
                return jsonify({'success': True})
                
            except Exception as e:
                logger.error(f"Error saving knowledge drafts: {str(e)}")
                return jsonify({'success': False, 'error': str(e)}), 500
            
        else:
            return jsonify({'success': False, 'error': 'Invalid action'}), 400
    
    # GET: Display the wizard
    return render_template('property_setup_wizard.html', user_id=user_id)


def search_local_places_google(property_address):
    """Search for local places using Gemini AI with Google Search"""
    try:
        from concierge.utils.gemini_config import gemini_model, genai_enabled
        
        if not genai_enabled or not gemini_model:
            logger.warning("Gemini AI not available, using fallback data")
            return get_fallback_places()
        
        # Use Gemini to search for local places
        search_prompt = f"""
        Find local recommendations near this address: {property_address}
        
        Please provide a list of places in the following categories, with real business names if possible:
        - Restaurants (3-4 recommendations)
        - Grocery stores/Supermarkets (2-3 recommendations)  
        - Coffee shops/Cafes (2-3 recommendations)
        - Medical facilities (Hospital, Urgent Care, Pharmacy) (2-3 recommendations)
        - Gas stations (1-2 recommendations)
        - Banks/ATMs (1-2 recommendations)
        - Tourist attractions/Points of interest (2-3 recommendations)
        - Parks/Recreation areas (1-2 recommendations)
        
        Format your response as a JSON array of objects with this structure:
        [
            {{"name": "Business Name", "type": "restaurant", "selected": false, "rating": 4.5, "address": "approximate address"}},
            ...
        ]
        
        Use these type categories: restaurant, shopping, cafe, medical, transportation, services, attractions, recreation
        
        Try to provide real, well-known businesses when possible. If you're not certain about specific businesses, use generic but realistic names like "Local Grocery Store" or "Downtown Coffee Shop".
        """
        
        try:
            response = gemini_model.generate_content(search_prompt)
            ai_response = response.text
            
            # Try to extract JSON from the response
            import json
            import re
            
            # Look for JSON array in the response
            json_match = re.search(r'\[.*\]', ai_response, re.DOTALL)
            if json_match:
                places_data = json.loads(json_match.group())
                
                # Validate and clean the data
                valid_places = []
                for place in places_data:
                    if isinstance(place, dict) and 'name' in place and 'type' in place:
                        valid_places.append({
                            'name': str(place.get('name', '')),
                            'type': str(place.get('type', 'general')),
                            'selected': False,
                            'rating': float(place.get('rating', 4.0)),
                            'address': str(place.get('address', ''))
                        })
                
                if valid_places:
                    logger.info(f"Found {len(valid_places)} places near {property_address} using Gemini AI")
                    return valid_places[:20]  # Limit to 20 places
                    
        except json.JSONDecodeError:
            logger.warning("Could not parse JSON from Gemini response, trying text parsing")
            # Fallback: try to parse text response
            places = parse_text_response_to_places(ai_response)
            if places:
                return places
                
        except Exception as e:
            logger.error(f"Error processing Gemini response: {str(e)}")
        
        # If AI search fails, return enhanced fallback data
        return get_enhanced_fallback_places(property_address)
        
    except Exception as e:
        logger.error(f"Error in Gemini-powered places search: {str(e)}")
        return get_fallback_places()


def parse_text_response_to_places(text_response):
    """Parse text response from Gemini into places data"""
    try:
        places = []
        lines = text_response.split('\n')
        
        current_type = 'general'
        type_mapping = {
            'restaurant': 'restaurant', 'cafe': 'cafe', 'coffee': 'cafe',
            'grocery': 'shopping', 'shopping': 'shopping', 'store': 'shopping',
            'medical': 'medical', 'hospital': 'medical', 'pharmacy': 'medical',
            'gas': 'transportation', 'bank': 'services', 'atm': 'services',
            'attraction': 'attractions', 'tourist': 'attractions',
            'park': 'recreation', 'recreation': 'recreation'
        }
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
                
            # Check if line indicates a new category
            for key, value in type_mapping.items():
                if key.lower() in line.lower() and ':' in line:
                    current_type = value
                    break
            
            # Extract business names (look for patterns like "- Name" or "1. Name")
            if re.match(r'^[-*•]\s*(.+)', line) or re.match(r'^\d+\.\s*(.+)', line):
                name_match = re.match(r'^[-*•\d.]\s*(.+)', line)
                if name_match:
                    name = name_match.group(1).strip()
                    places.append({
                        'name': name,
                        'type': current_type,
                        'selected': False,
                        'rating': 4.0 + (len(places) % 10) * 0.1,  # Vary ratings slightly
                        'address': ''
                    })
        
        return places[:20] if places else []
        
    except Exception as e:
        logger.error(f"Error parsing text response: {str(e)}")
        return []


def get_enhanced_fallback_places(property_address):
    """Enhanced fallback places with address-specific suggestions"""
    base_places = get_fallback_places()
    
    # Add some address-specific context if possible
    city = "Local"
    if property_address:
        parts = property_address.split(',')
        if len(parts) > 1:
            city = parts[-2].strip() if len(parts) > 2 else parts[-1].strip()
    
    # Enhance names with city context
    enhanced_places = []
    for place in base_places:
        enhanced_place = place.copy()
        if 'Local' in place['name']:
            enhanced_place['name'] = place['name'].replace('Local', city)
        enhanced_places.append(enhanced_place)
    
    return enhanced_places


def process_uploaded_file(file, upload_type):
    """Process uploaded files for the wizard"""
    import tempfile
    import json
    import pandas as pd
    from werkzeug.utils import secure_filename
    
    filename = secure_filename(file.filename)
    file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    try:
        # Save file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_ext}') as temp_file:
            file.save(temp_file.name)
            
            if upload_type == 'json' and file_ext == 'json':
                # Process Airbnb JSON file
                return process_airbnb_json(temp_file.name)
            
            elif upload_type == 'facts':
                # Process property facts file
                return process_property_facts_file(temp_file.name, file_ext)
            
            else:
                return {'error': f'Unsupported file type: {file_ext}'}
                
    except Exception as e:
        logger.error(f"Error processing uploaded file {filename}: {str(e)}")
        return {'error': f'Failed to process file: {str(e)}'}
    finally:
        # Clean up temp file
        try:
            import os
            if 'temp_file' in locals():
                os.unlink(temp_file.name)
        except:
            pass


def process_airbnb_json(file_path):
    """Process Airbnb JSON data file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Extract relevant information from Airbnb JSON
        extracted_data = {}
        
        # Try to extract property info, amenities, house rules, etc.
        if isinstance(data, dict):
            # Look for common Airbnb export fields
            if 'listing' in data:
                listing = data['listing']
                extracted_data['name'] = listing.get('name', '')
                extracted_data['description'] = listing.get('description', '')
                extracted_data['house_rules'] = listing.get('house_rules', '')
                extracted_data['amenities'] = listing.get('amenities', [])
                
            # Look for calendar/reservation data
            if 'calendar' in data:
                extracted_data['calendar_data'] = len(data['calendar'])
                
        return {
            'type': 'airbnb_data',
            'data': extracted_data,
            'message': 'Airbnb data processed successfully'
        }
        
    except Exception as e:
        return {'error': f'Failed to process Airbnb JSON: {str(e)}'}


def process_property_facts_file(file_path, file_ext):
    """Process property facts from various file types"""
    try:
        content = ''
        
        if file_ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
        elif file_ext == 'pdf':
            # Process PDF file
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = ''
                    for page in reader.pages:
                        content += page.extract_text() + '\n'
            except ImportError:
                return {'error': 'PDF processing not available (PyPDF2 not installed)'}
                
        elif file_ext in ['xlsx', 'xls']:
            # Process Excel file
            try:
                df = pd.read_excel(file_path)
                content = df.to_string()
            except ImportError:
                return {'error': 'Excel processing not available (pandas not installed)'}
                
        elif file_ext == 'docx':
            # Process Word document
            try:
                from docx import Document
                doc = Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                return {'error': 'Word document processing not available (python-docx not installed)'}
                
        elif file_ext == 'json':
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                content = json.dumps(data, indent=2)
        
        else:
            return {'error': f'Unsupported file type: {file_ext}'}
        
        # Extract structured information from content using AI if available
        if genai_enabled and content:
            structured_data = extract_property_facts_with_ai(content)
        else:
            structured_data = {'raw_content': content}
        
        return {
            'type': 'property_facts',
            'data': structured_data,
            'message': f'Property facts extracted from {file_ext.upper()} file'
        }
        
    except Exception as e:
        return {'error': f'Failed to process property facts file: {str(e)}'}


def extract_property_facts_with_ai(content):
    """Use Gemini to extract structured property facts from content"""
    try:
        if not gemini_model:
            return {'raw_content': content}
        
        prompt = f"""
        Extract structured property information from the following content. 
        Return the information organized by categories like:
        - Rooms and Layout (bedrooms, bathrooms, living areas)
        - Amenities and Features (kitchen appliances, electronics, etc.)
        - House Rules and Policies
        - Check-in/Check-out Information
        - WiFi and Technical Details
        - Emergency Information
        - Local Area Information
        
        Content:
        {content[:3000]}  # Limit content to avoid token limits
        
        Return a JSON-like structure with clear categories and details.
        """
        
        response = gemini_model.generate_content(prompt)
        
        # Try to parse AI response as structured data
        ai_response = response.text
        
        return {
            'raw_content': content,
            'ai_extracted': ai_response,
            'extraction_timestamp': datetime.now(timezone.utc).isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error extracting property facts with AI: {str(e)}")
        return {
            'raw_content': content,
            'ai_extraction_error': str(e)
        }


def get_fallback_places():
    """Fallback places data when Google API is not available - host favorites only"""
    return [
        {
            'name': 'Girl & the Goat', 
            'type': 'restaurant', 
            'selected': False, 
            'rating': 4.5,
            'address': '809 W Randolph St, Chicago, IL 60607',
            'description': 'Innovative Mediterranean cuisine in a trendy setting'
        },
        {
            'name': 'Starbucks Reserve Roastery', 
            'type': 'cafe', 
            'selected': False, 
            'rating': 4.3,
            'address': '646 N Michigan Ave, Chicago, IL 60611',
            'description': 'Premium coffee experience with artisanal roasts'
        },
        {
            'name': 'Millennium Park', 
            'type': 'recreation', 
            'selected': False, 
            'rating': 4.4,
            'address': '201 E Randolph St, Chicago, IL 60602',
            'description': 'Iconic park featuring Cloud Gate and outdoor concerts'
        },
        {
            'name': 'Art Institute of Chicago', 
            'type': 'attractions', 
            'selected': False, 
            'rating': 4.6,
            'address': '111 S Michigan Ave, Chicago, IL 60603',
            'description': 'World-renowned art museum with impressive collections'
        },
        {
            'name': 'Navy Pier', 
            'type': 'attractions', 
            'selected': False, 
            'rating': 4.2,
            'address': '600 E Grand Ave, Chicago, IL 60611',
            'description': 'Waterfront entertainment destination with shops and dining'
        },
        {
            'name': 'The Purple Pig', 
            'type': 'restaurant', 
            'selected': False, 
            'rating': 4.4,
            'address': '500 N Michigan Ave, Chicago, IL 60611',
            'description': 'Mediterranean small plates and natural wines'
        },
        {
            'name': 'Lincoln Park Zoo', 
            'type': 'attractions', 
            'selected': False, 
            'rating': 4.3,
            'address': '2001 N Clark St, Chicago, IL 60614',
            'description': 'Free admission zoo in beautiful Lincoln Park'
        },
        {
            'name': 'Intelligentsia Coffee', 
            'type': 'cafe', 
            'selected': False, 
            'rating': 4.4,
            'address': '53 W Jackson Blvd, Chicago, IL 60604',
            'description': 'Specialty coffee roasters with expertly crafted drinks'
        },
        {
            'name': 'Whole Foods Market', 
            'type': 'shopping', 
            'selected': False, 
            'rating': 4.1,
            'address': '1550 N Kingsbury St, Chicago, IL 60642',
            'description': 'Organic grocery store with prepared foods and local products'
        }
    ]


def process_wizard_knowledge(property_id, wizard_data):
    """Process wizard data into knowledge items"""
    try:
        from concierge.utils.firestore_client import create_knowledge_item
        import uuid
        knowledge_items_created = 0
        
        # Process house rules
        house_rules = wizard_data.get('house_rules', {})
        rules_list = []
        
        # Standard rules that are enabled
        for rule_key, rule_data in house_rules.items():
            if isinstance(rule_data, dict) and rule_data.get('enabled'):
                rules_list.append(rule_data.get('text', ''))
        
        # Custom rules
        custom_rules = house_rules.get('custom_rules', [])
        for rule in custom_rules:
            if rule.get('text'):
                rules_list.append(rule['text'])
        
        if rules_list:
            rules_content = "House Rules:\n" + "\n".join([f"• {rule}" for rule in rules_list])
            # Generate a unique ID for the knowledge item
            item_id = str(uuid.uuid4())
            knowledge_item_data = {
                'propertyId': property_id,
                'type': 'house_rules',
                'content': rules_content,
                'tags': ['house_rules', 'rules', 'property_info'],
                'status': 'approved',
                'source': 'setup_wizard',
                'created_at': datetime.now(timezone.utc)
            }
            if create_knowledge_item(item_id, knowledge_item_data):
                knowledge_items_created += 1
        
        # Process emergency information
        emergency_info = wizard_data.get('emergency_info', {})
        emergency_list = []
        
        # Standard emergencies that are enabled
        for emergency_key, emergency_data in emergency_info.items():
            if isinstance(emergency_data, dict) and emergency_data.get('enabled'):
                title = emergency_data.get('title', emergency_key.replace('_', ' ').title())
                instructions = emergency_data.get('instructions', '')
                location = emergency_data.get('location', '')
                
                emergency_text = f"{title}"
                if instructions:
                    emergency_text += f"\nInstructions: {instructions}"
                if location:
                    emergency_text += f"\nLocation: {location}"
                
                emergency_list.append(emergency_text)
        
        # Custom emergencies
        custom_emergencies = emergency_info.get('custom_emergencies', [])
        for emergency in custom_emergencies:
            if emergency.get('title'):
                emergency_text = emergency['title']
                if emergency.get('instructions'):
                    emergency_text += f"\nInstructions: {emergency['instructions']}"
                if emergency.get('location'):
                    emergency_text += f"\nLocation: {emergency['location']}"
                emergency_list.append(emergency_text)
        
        if emergency_list:
            emergency_content = "Emergency Information:\n" + "\n\n".join([f"• {emergency}" for emergency in emergency_list])
            # Generate a unique ID for the knowledge item
            item_id = str(uuid.uuid4())
            knowledge_item_data = {
                'propertyId': property_id,
                'type': 'emergency_info',
                'content': emergency_content,
                'tags': ['emergency', 'safety', 'property_info'],
                'status': 'approved',
                'source': 'setup_wizard',
                'created_at': datetime.now(timezone.utc)
            }
            if create_knowledge_item(item_id, knowledge_item_data):
                knowledge_items_created += 1
        
        # Process local recommendations
        local_places = wizard_data.get('local_recommendations', {}).get('places', [])
        selected_places = [place for place in local_places if place.get('selected')]
        
        if selected_places:
            places_by_type = {}
            for place in selected_places:
                place_type = place.get('type', 'general')
                if place_type not in places_by_type:
                    places_by_type[place_type] = []
                places_by_type[place_type].append(place.get('name', ''))
            
            recommendations_content = "Local Recommendations:\n"
            for place_type, places in places_by_type.items():
                recommendations_content += f"\n{place_type.title()}:\n"
                recommendations_content += "\n".join([f"• {place}" for place in places])
            
            # Generate a unique ID for the knowledge item
            item_id = str(uuid.uuid4())
            knowledge_item_data = {
                'propertyId': property_id,
                'type': 'local_recommendations',
                'content': recommendations_content,
                'tags': ['local', 'recommendations', 'area_info'],
                'status': 'approved',
                'source': 'setup_wizard',
                'created_at': datetime.now(timezone.utc)
            }
            if create_knowledge_item(item_id, knowledge_item_data):
                knowledge_items_created += 1
        
        # Process basic property information
        basic_info = wizard_data.get('basic_info', {})
        if basic_info:
            property_info_parts = []
            
            if basic_info.get('description'):
                property_info_parts.append(f"Property Description: {basic_info['description']}")
            
            if basic_info.get('checkin_time') or basic_info.get('checkout_time'):
                checkin = basic_info.get('checkin_time', '15:00')
                checkout = basic_info.get('checkout_time', '11:00')
                property_info_parts.append(f"Check-in: {checkin}, Check-out: {checkout}")
            
            if basic_info.get('wifi_network'):
                wifi_info = f"WiFi Network: {basic_info['wifi_network']}"
                if basic_info.get('wifi_password'):
                    wifi_info += f"\nWiFi Password: {basic_info['wifi_password']}"
                property_info_parts.append(wifi_info)
            
            if property_info_parts:
                property_content = "Property Information:\n" + "\n\n".join(property_info_parts)
                # Generate a unique ID for the knowledge item
                item_id = str(uuid.uuid4())
                knowledge_item_data = {
                    'propertyId': property_id,
                    'type': 'property_info',
                    'content': property_content,
                    'tags': ['property', 'basic_info', 'checkin', 'wifi'],
                    'status': 'approved',
                    'source': 'setup_wizard',
                    'created_at': datetime.now(timezone.utc)
                }
                if create_knowledge_item(item_id, knowledge_item_data):
                    knowledge_items_created += 1
        
        return knowledge_items_created
        
    except Exception as e:
        logger.error(f"Error processing wizard knowledge for property {property_id}: {str(e)}")
        return 0


def process_approved_knowledge_drafts(property_id, wizard_data):
    """Process approved knowledge drafts into actual knowledge items"""
    try:
        from concierge.utils.firestore_client import create_knowledge_item
        import uuid
        knowledge_items_created = 0
        
        # Get approved knowledge drafts from wizard data
        knowledge_drafts = wizard_data.get('knowledge_drafts', [])
        approved_drafts = [draft for draft in knowledge_drafts if draft.get('status') == 'approved']
        
        logger.info(f"Processing {len(approved_drafts)} approved knowledge drafts for property {property_id}")
        
        for draft in approved_drafts:
            try:
                # Generate unique ID for the knowledge item
                item_id = str(uuid.uuid4())
                
                # Prepare knowledge item data following the unified schema
                knowledge_item_data = {
                    'propertyId': property_id,
                    'type': draft.get('type', 'other'),
                    'content': draft.get('content', ''),
                    'tags': draft.get('tags', []),
                    'status': 'approved',
                    'source': 'setup_wizard',
                    'created_at': datetime.now(timezone.utc),
                    'updated_at': datetime.now(timezone.utc)
                }
                
                # Add any additional fields from the draft
                if 'question' in draft:
                    knowledge_item_data['question'] = draft['question']
                if 'category' in draft:
                    knowledge_item_data['category'] = draft['category']
                
                # Create the knowledge item in Firestore
                success = create_knowledge_item(item_id, knowledge_item_data)
                if success:
                    knowledge_items_created += 1
                    logger.info(f"Created knowledge item {item_id} from draft {draft.get('id', 'unknown')}")
                else:
                    logger.warning(f"Failed to create knowledge item from draft {draft.get('id', 'unknown')}")
                    
            except Exception as draft_error:
                logger.error(f"Error processing individual draft {draft.get('id', 'unknown')}: {str(draft_error)}")
                continue
        
        # Also process voice session data if available and not already in drafts
        voice_sessions = wizard_data.get('voice_sessions', [])
        if voice_sessions and knowledge_items_created == 0:
            logger.info("No approved drafts found, processing voice session data as fallback")
            
            for session in voice_sessions:
                voice_collected_data = session.get('voiceCollectedData', [])
                for voice_item in voice_collected_data:
                    try:
                        item_id = str(uuid.uuid4())
                        
                        # Convert voice data to knowledge item
                        content = voice_item.get('fullResponse', voice_item.get('instruction', ''))
                        topic = voice_item.get('topic', 'other')
                        
                        # Map voice topics to knowledge types
                        type_mapping = {
                            'wifi': 'connectivity',
                            'parking': 'access_info', 
                            'emergency': 'emergency_info',
                            'appliances': 'amenities'
                        }
                        
                        knowledge_item_data = {
                            'propertyId': property_id,
                            'type': type_mapping.get(topic, 'other'),
                            'content': content,
                            'tags': [topic, 'voice_session', 'setup_wizard'],
                            'status': 'approved',
                            'source': 'voice_session',
                            'created_at': datetime.now(timezone.utc),
                            'updated_at': datetime.now(timezone.utc)
                        }
                        
                        success = create_knowledge_item(item_id, knowledge_item_data)
                        if success:
                            knowledge_items_created += 1
                            
                    except Exception as voice_error:
                        logger.error(f"Error processing voice data: {str(voice_error)}")
                        continue
        
        logger.info(f"Successfully created {knowledge_items_created} knowledge items for property {property_id}")
        return knowledge_items_created
        
    except Exception as e:
        logger.error(f"Error processing approved knowledge drafts for property {property_id}: {str(e)}")
        return 0


def generate_wizard_knowledge_drafts(wizard_data):
    """Generate knowledge drafts using Gemini AI from all wizard data"""
    try:
        from concierge.api.utils import generate_knowledge_items_with_gemini
        import uuid
        
        logger.info(f"Starting knowledge generation for wizard data with {len(wizard_data)} sections")
        
        # Extract property details
        basic_info = wizard_data.get('basic_info', {})
        property_details = {
            'name': basic_info.get('property_name', 'Property'),
            'address': basic_info.get('property_address', ''),
            'id': 'wizard_temp'
        }
        
        logger.info(f"Property details: {property_details}")
        
        # Compile all wizard data into a comprehensive text for AI processing
        wizard_text = compile_wizard_data_for_ai(wizard_data)
        
        if not wizard_text or len(wizard_text.strip()) < 50:
            logger.warning(f"Insufficient wizard data to process: {len(wizard_text)} characters")
            return []
        
        logger.info(f"Calling AI generation with {len(wizard_text)} characters of text")
        
        # Use existing knowledge generation function
        generated_items = generate_knowledge_items_with_gemini(wizard_text, property_details)
        
        logger.info(f"AI returned {len(generated_items)} generated items")
        
        # Convert to draft format with unique IDs and pending status
        knowledge_drafts = []
        for item in generated_items:
            # Extract title from content if not present
            title = item.get('title')
            if not title:
                content_lines = item.get('content', '').strip().split('\n')
                title = content_lines[0][:100] if content_lines else 'Knowledge Item'
            
            # Get category from type
            category_map = {
                'house_rules': 'House Rules',
                'emergency_info': 'Emergency Information',  
                'property_info': 'Basic Information',
                'local_recommendations': 'Local Recommendations',
                'voice_data': 'Voice Session Data',
                'property_facts': 'Property Facts'
            }
            
            draft = {
                'id': str(uuid.uuid4()),
                'type': item.get('type', 'other'),
                'title': title,
                'content': item.get('content', ''),
                'tags': item.get('tags', []),
                'category': category_map.get(item.get('type'), 'General'),
                'status': 'pending',
                'source': 'ai_generated',
                'created_at': datetime.now(timezone.utc).isoformat()
            }
            knowledge_drafts.append(draft)
        
        logger.info(f"Generated {len(knowledge_drafts)} knowledge drafts from wizard data")
        return knowledge_drafts
        
    except Exception as e:
        logger.error(f"Error generating knowledge drafts: {str(e)}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return []


def compile_wizard_data_for_ai(wizard_data):
    """Compile all wizard data into a comprehensive text for AI processing"""
    text_parts = []
    
    logger.info(f"Compiling wizard data with keys: {list(wizard_data.keys())}")
    
    # Basic Information
    basic_info = wizard_data.get('basic_info', {})
    if basic_info:
        text_parts.append("=== BASIC PROPERTY INFORMATION ===")
        if basic_info.get('property_name'):
            text_parts.append(f"Property Name: {basic_info['property_name']}")
        if basic_info.get('property_address'):
            text_parts.append(f"Address: {basic_info['property_address']}")
        if basic_info.get('description'):
            text_parts.append(f"Description: {basic_info['description']}")
        if basic_info.get('checkin_time'):
            text_parts.append(f"Check-in Time: {basic_info['checkin_time']}")
        if basic_info.get('checkout_time'):
            text_parts.append(f"Check-out Time: {basic_info['checkout_time']}")
        if basic_info.get('wifi_network'):
            text_parts.append(f"WiFi Network: {basic_info['wifi_network']}")
        if basic_info.get('wifi_password'):
            text_parts.append(f"WiFi Password: {basic_info['wifi_password']}")
        text_parts.append("")
    
    # House Rules
    house_rules = wizard_data.get('house_rules', {})
    if house_rules:
        text_parts.append("=== HOUSE RULES ===")
        
        # Handle both object and array formats
        if isinstance(house_rules, list):
            # Array format from frontend
            for rule in house_rules:
                if isinstance(rule, dict) and rule.get('enabled'):
                    rule_text = rule.get('text', '')
                    if rule_text:
                        text_parts.append(rule_text)
        else:
            # Object format
            # Standard rules
            standard_rules = []
            for rule_key, rule_data in house_rules.items():
                if isinstance(rule_data, dict) and rule_data.get('enabled'):
                    rule_text = rule_data.get('text', rule_key.replace('_', ' ').title())
                    standard_rules.append(rule_text)
            
            if standard_rules:
                text_parts.extend(standard_rules)
            
            # Custom rules
            custom_rules = house_rules.get('custom_rules', [])
            for rule in custom_rules:
                if rule.get('text'):
                    text_parts.append(rule['text'])
        
        text_parts.append("")
    
    # Emergency Information
    emergency_info = wizard_data.get('emergency_info', {})
    if emergency_info:
        text_parts.append("=== EMERGENCY INFORMATION ===")
        
        # Handle both object and array formats
        if isinstance(emergency_info, list):
            # Array format from frontend
            for emergency in emergency_info:
                if isinstance(emergency, dict) and emergency.get('enabled'):
                    title = emergency.get('title', emergency.get('scenario', ''))
                    instructions = emergency.get('instructions', '')
                    location = emergency.get('location', '')
                    
                    if title:
                        text_parts.append(f"{title}:")
                        if instructions:
                            text_parts.append(f"  Instructions: {instructions}")
                        if location:
                            text_parts.append(f"  Location: {location}")
        else:
            # Object format
            # Standard emergencies
            for emergency_key, emergency_data in emergency_info.items():
                if isinstance(emergency_data, dict) and emergency_data.get('enabled'):
                    title = emergency_data.get('title', emergency_key.replace('_', ' ').title())
                    instructions = emergency_data.get('instructions', '')
                    location = emergency_data.get('location', '')
                    
                    text_parts.append(f"{title}:")
                    if instructions:
                        text_parts.append(f"  Instructions: {instructions}")
                    if location:
                        text_parts.append(f"  Location: {location}")
            
            # Custom emergencies
            custom_emergencies = emergency_info.get('custom_emergencies', [])
            for emergency in custom_emergencies:
                if emergency.get('title'):
                    text_parts.append(f"{emergency['title']}:")
                    if emergency.get('instructions'):
                        text_parts.append(f"  Instructions: {emergency['instructions']}")
                    if emergency.get('location'):
                        text_parts.append(f"  Location: {emergency['location']}")
        
        text_parts.append("")
    
    # Local Recommendations
    local_recommendations = wizard_data.get('local_recommendations', {})
    if local_recommendations:
        text_parts.append("=== LOCAL RECOMMENDATIONS ===")
        places = local_recommendations.get('places', [])
        selected_places = [place for place in places if place.get('selected')]
        
        if selected_places:
            places_by_type = {}
            for place in selected_places:
                place_type = place.get('type', 'general')
                if place_type not in places_by_type:
                    places_by_type[place_type] = []
                
                place_info = place.get('name', '')
                if place.get('address'):
                    place_info += f" - {place['address']}"
                if place.get('description'):
                    place_info += f" ({place['description']})"
                
                places_by_type[place_type].append(place_info)
            
            for place_type, places in places_by_type.items():
                text_parts.append(f"{place_type.title()}:")
                for place in places:
                    text_parts.append(f"  - {place}")
        
        text_parts.append("")
    
    # Property Facts (files, manual entry, voice data)
    property_facts = wizard_data.get('property_facts', {})
    if property_facts:
        text_parts.append("=== PROPERTY FACTS ===")
        
        # File extracts
        if property_facts.get('extracted_text'):
            text_parts.append("From uploaded files:")
            text_parts.append(property_facts['extracted_text'])
        
        # Manual entry
        manual_entry = wizard_data.get('manual_entry', {})
        if manual_entry and manual_entry.get('rooms'):
            text_parts.append("Manual room configuration:")
            for room in manual_entry['rooms']:
                room_name = room.get('name', 'Room')
                room_type = room.get('type', 'room')
                text_parts.append(f"  {room_name} ({room_type}):")
                
                amenities = room.get('amenities', [])
                if amenities:
                    text_parts.append(f"    Amenities: {', '.join(amenities)}")
                
                if room.get('notes'):
                    text_parts.append(f"    Notes: {room['notes']}")
        
        text_parts.append("")
    
    # Voice Session Data
    voice_sessions = wizard_data.get('voice_sessions', [])
    if voice_sessions:
        text_parts.append("=== VOICE SESSION INFORMATION ===")
        
        for session in voice_sessions:
            if session.get('extractedAnswers'):
                text_parts.append(f"From voice session {session.get('sessionId', 'unknown')[:8]}:")
                
                for topic, answer in session['extractedAnswers'].items():
                    if topic == 'wifi' and isinstance(answer, dict):
                        if answer.get('networkName') and answer.get('password'):
                            text_parts.append(f"  WiFi: Network '{answer['networkName']}', Password '{answer['password']}'")
                    elif topic == 'parking' and isinstance(answer, dict):
                        if answer.get('instruction'):
                            text_parts.append(f"  Parking: {answer['instruction']}")
                    elif topic == 'emergency' and isinstance(answer, dict):
                        if answer.get('detail'):
                            text_parts.append(f"  Emergency: {answer['detail']}")
                    elif topic == 'appliances' and isinstance(answer, dict):
                        if answer.get('brandModel'):
                            text_parts.append(f"  Appliance: {answer['brandModel']}")
        
        text_parts.append("")
    
    # Airbnb/Calendar Data
    airbnb_data = wizard_data.get('airbnb_data', {})
    if airbnb_data:
        text_parts.append("=== AIRBNB/CALENDAR DATA ===")
        
        if airbnb_data.get('ical_url'):
            text_parts.append(f"Calendar URL: {airbnb_data['ical_url']}")
        
        if airbnb_data.get('json_data'):
            text_parts.append("Property listing data:")
            # Try to extract meaningful info from JSON
            try:
                import json
                if isinstance(airbnb_data['json_data'], str):
                    json_data = json.loads(airbnb_data['json_data'])
                else:
                    json_data = airbnb_data['json_data']
                
                # Extract common fields
                if isinstance(json_data, dict):
                    if json_data.get('name'):
                        text_parts.append(f"  Listing Name: {json_data['name']}")
                    if json_data.get('description'):
                        text_parts.append(f"  Description: {json_data['description']}")
                    if json_data.get('amenities'):
                        text_parts.append(f"  Amenities: {', '.join(json_data['amenities'])}")
                    if json_data.get('house_rules'):
                        text_parts.append(f"  House Rules: {json_data['house_rules']}")
                        
            except Exception as e:
                text_parts.append(f"  Raw data: {str(airbnb_data['json_data'])[:500]}...")
        
        text_parts.append("")
    
    compiled_text = "\n".join(text_parts)
    logger.info(f"Compiled wizard data: {len(compiled_text)} characters, {len(text_parts)} sections")
    if len(compiled_text) < 100:
        logger.warning(f"Very short compiled text: {compiled_text}")
    
    return compiled_text